# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"C:\Users\user\Documents\Credit card web project")
OUT_DIR = ROOT / "outputs" / "product-spec"
ASSET_DIR = OUT_DIR / "assets-v20260603-zh"
DOCX_PATH = OUT_DIR / "credit-card-mvp-spec-2026-06-03-v3.docx"
TEXT_DUMP = OUT_DIR / "credit-card-mvp-spec-2026-06-03-v3-inspection.txt"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

PALETTE = {
    "ink": (33, 43, 54),
    "navy": (20, 45, 78),
    "teal": (46, 122, 142),
    "mist": (240, 245, 247),
    "line": (207, 215, 222),
    "gold": (201, 161, 77),
    "gray": (95, 104, 112),
    "bg": (250, 246, 238),
    "white": (255, 255, 255),
}


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msjhbd.ttc" if bold else "C:/Windows/Fonts/msjh.ttc"),
        Path("C:/Windows/Fonts/mingliub.ttc" if bold else "C:/Windows/Fonts/mingliu.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


F_TITLE = font(42, True)
F_H1 = font(30, True)
F_H2 = font(22, True)
F_BODY = font(18, False)
F_SMALL = font(16, False)


def rounded(draw, box, fill, outline=None, radius=18, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def title_block(draw, title, subtitle):
    draw.text((60, 42), title, font=F_TITLE, fill=PALETTE["navy"])
    draw.text((60, 102), subtitle, font=F_BODY, fill=PALETTE["ink"])
    draw.line((60, 148, 1540, 148), fill=PALETTE["line"], width=3)


def create_category_rule_image(path: Path):
    img = Image.new("RGB", (1600, 980), PALETTE["white"])
    draw = ImageDraw.Draw(img)
    title_block(draw, "分類頁卡片規則：固定高度與二行摘要", "列表頁負責快速比較，因此推薦理由只顯示二行；超過內容由詳情頁承接，不在列表頁無限展開。")
    draw.text((60, 188), "現金回饋", font=F_H1, fill=PALETTE["navy"])
    draw.text((60, 228), "卡片高度需可控，避免少數長內容把後面卡片整片往下擠。", font=F_BODY, fill=PALETTE["ink"])
    rounded(draw, (60, 278, 1540, 352), PALETTE["mist"], PALETTE["line"])
    filters = ["銀行：全部", "僅看進行中", "排序：本站推薦優先", "摘要：固定二行"]
    x = 84
    for item in filters:
        rounded(draw, (x, 298, x + 260, 334), PALETTE["white"], radius=10)
        draw.text((x + 16, 306), item, font=F_SMALL, fill=PALETTE["ink"])
        x += 280
    rows = [
        ("1", "台新太陽卡", "台新銀行", "一般消費最高 5%", "推薦理由：適合日常消費，門檻低、規則清楚。", "若還有更多限制條件，改在詳情頁完整說明。", True),
        ("2", "國泰 CUBE 卡", "國泰世華", "指定通路最高 3%", "推薦理由：適合數位消費，主題切換邏輯明確。", "列表頁只保留比較所需重點，不延伸長文。", False),
        ("3", "富邦 J 卡", "台北富邦", "旅遊場景最高 5%", "推薦理由：旅日場景鮮明，適合作為主題整理卡片。", "完整活動細節與限制條件交由詳情頁承接。", False),
    ]
    y = 388
    for rank, title, bank, reward, line1, line2, featured in rows:
        rounded(draw, (60, y, 1540, y + 148), PALETTE["white"], PALETTE["line"], radius=20)
        rounded(draw, (82, y + 28, 138, y + 84), PALETTE["navy"], radius=16)
        draw.text((104, y + 35), rank, font=F_H2, fill=PALETTE["white"])
        rounded(draw, (166, y + 20, 314, y + 118), (229, 236, 241), radius=16)
        draw.text((208, y + 56), "卡面圖", font=F_H2, fill=PALETTE["navy"])
        draw.text((344, y + 22), title, font=F_H2, fill=PALETTE["ink"])
        draw.text((344, y + 52), bank, font=F_SMALL, fill=PALETTE["gray"])
        draw.text((344, y + 82), line1, font=F_SMALL, fill=PALETTE["gray"])
        draw.text((344, y + 108), line2, font=F_SMALL, fill=PALETTE["gray"])
        draw.text((1060, y + 28), reward, font=F_H1, fill=PALETTE["teal"])
        draw.text((1060, y + 74), "優惠期間：2026/06/01 - 2026/12/31", font=F_SMALL, fill=PALETTE["ink"])
        if featured:
            rounded(draw, (1328, y + 28, 1498, y + 64), PALETTE["gold"], radius=16)
            draw.text((1354, y + 36), "本站推薦", font=F_SMALL, fill=PALETTE["ink"])
        y += 170
    img.save(path)


def create_admin_editor_image(path: Path):
    img = Image.new("RGB", (1600, 1120), PALETTE["white"])
    draw = ImageDraw.Draw(img)
    title_block(draw, "後台畫面補強：單筆優惠編輯與排序設定", "第 7 與第 8 節不只要列出欄位名稱，還要讓閱讀者看懂後台如何填寫，以及如何影響前台排序與呈現。")
    rounded(draw, (60, 184, 300, 1060), PALETTE["navy"], radius=26)
    menu = ["儀表板", "銀行管理", "信用卡管理", "分類管理", "優惠管理", "設定"]
    y = 238
    for item in menu:
        rounded(draw, (82, y, 278, y + 48), PALETTE["white"], radius=14)
        draw.text((110, y + 11), item, font=F_BODY, fill=PALETTE["navy"])
        y += 74
    rounded(draw, (330, 184, 1540, 1060), PALETTE["white"], PALETTE["line"], radius=24)
    draw.text((364, 220), "優惠編輯：CUBE 指定通路回饋", font=F_H1, fill=PALETTE["navy"])
    sections = [
        ("基本資訊", [
            ("優惠標題", "文字輸入", "控制前台卡片標題與詳情頁主標"),
            ("所屬分類", "下拉選單", "決定分類頁歸屬"),
            ("關聯信用卡", "多選欄位", "支援一個優惠掛多張卡"),
        ]),
        ("列表摘要設定", [
            ("summary_mode", "單選按鈕", "系統摘要或人工摘要"),
            ("target_audience", "下拉選單", "用於組成推薦理由中的適合族群"),
            ("highlight_1 / highlight_2", "下拉選單", "用於組成推薦理由中的主優勢"),
            ("manual_summary", "文字輸入（選填）", "若有值，前台優先顯示人工摘要"),
            ("summary_preview", "唯讀預覽", "預覽前台列表卡片實際看到的二行摘要"),
        ]),
        ("排序與上架設定", [
            ("is_featured", "開關", "控制本站推薦標記，並優先排序"),
            ("recommend_score", "數字輸入 0-100", "分數越高越前面"),
            ("sort_order", "整數輸入", "同分時人工微調，數字越小越前"),
            ("is_published", "開關", "控制是否可出現在前台"),
            ("updated_at", "系統自動帶入", "其他排序條件相同時，作為最後排序依據"),
        ]),
    ]
    y = 280
    for sec_title, rows in sections:
        height = 64 + len(rows) * 56
        rounded(draw, (364, y, 1508, y + height), PALETTE["mist"], PALETTE["line"], radius=18)
        draw.text((390, y + 18), sec_title, font=F_H2, fill=PALETTE["navy"])
        ry = y + 58
        for label, input_type, desc in rows:
            draw.line((390, ry, 1480, ry), fill=PALETTE["line"], width=1)
            draw.text((398, ry + 12), label, font=F_SMALL, fill=PALETTE["ink"])
            draw.text((702, ry + 12), input_type, font=F_SMALL, fill=PALETTE["gray"])
            draw.text((980, ry + 12), desc, font=F_SMALL, fill=PALETTE["gray"])
            ry += 56
        y = ry + 22
    img.save(path)


def set_fill(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def style(run, size=10.5, bold=False, color="212B36"):
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def heading(doc, text, intro=None):
    p = doc.add_paragraph()
    style(p.add_run(text), 18, True, "14334E")
    if intro:
        p2 = doc.add_paragraph()
        style(p2.add_run(intro), 10.5, False, "394651")


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style(p.add_run(item), 10.5)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        style(p.add_run(item), 10.5)


def kv_table(doc, rows, col1=2.0, col2=4.3):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(col1)
    table.columns[1].width = Inches(col2)
    for left, right in rows:
        cells = table.add_row().cells
        set_fill(cells[0], "F0F5F7")
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        style(cells[0].paragraphs[0].add_run(left), 10, True, "14334E")
        style(cells[1].paragraphs[0].add_run(right), 10, False, "212B36")
    doc.add_paragraph()


def grid_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)
    for idx, header in enumerate(headers):
        set_fill(table.cell(0, idx), "DDE7ED")
        style(table.cell(0, idx).paragraphs[0].add_run(header), 9.5, True, "14334E")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            style(cells[idx].paragraphs[0].add_run(value), 9.3, False, "212B36")
    doc.add_paragraph()


category_img = ASSET_DIR / "分類頁卡片規則-2026-06-03.png"
admin_img = ASSET_DIR / "後台單筆優惠編輯畫面-2026-06-03.png"
create_category_rule_image(category_img)
create_admin_editor_image(admin_img)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
normal = doc.styles["Normal"]
normal.font.name = "Microsoft JhengHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
normal.font.size = Pt(10.5)

p = doc.add_paragraph()
style(p.add_run("信用卡優惠查詢網站 MVP 規格書"), 24, True, "14334E")
p2 = doc.add_paragraph()
style(p2.add_run("版本：2026-06-03 / v3\n本版依 discussion-checklist.md 補強分類列表頁規則、第 7 節資料結構規劃與第 8 節商業規則與排序，並修正上一版出現的亂碼問題。"), 11, False, "394651")

heading(doc, "1. 產品目標", "這個網站不是冷資料庫，而是讓使用者可以快速理解哪張卡值得看、為什麼值得看的整理型入口站。")
bullets(doc, [
    "前台提供公開查詢，使用者可從分類、關鍵字、銀行與是否進行中快速篩選。",
    "首頁偏導覽與整理感，透過圖文卡片、分類入口與精選區塊降低閱讀負擔。",
    "後台優先解決資料治理，讓管理者能維護銀行、信用卡、優惠、分類與推薦排序。",
])

heading(doc, "2. MVP 範圍", "第一版聚焦在能上線、能維護、能擴充，而不是一次做到最完整。")
numbered(doc, [
    "前台：首頁、分類列表頁、優惠詳情頁、搜尋與篩選。",
    "後台：單一管理員登入，銀行 / 卡片 / 分類 / 優惠資料 CRUD。",
    "排序：採手動推薦權重，後台可編輯 is_featured、recommend_score、sort_order。",
    "資料：先人工建檔，預留未來接銀行官網爬蟲與規則評分。",
])

heading(doc, "3. 首頁概念", "首頁是懶人包式入口頁。左側讓使用者快速切入優惠主題，右側集中展示最新與值得先看的信用卡。")
kv_table(doc, [
    ("核心任務", "讓使用者 10 秒內找到自己想看的優惠主題。"),
    ("主要元件", "分類導覽、搜尋框、精選卡片、進行中與推薦標籤。"),
    ("視覺方向", "資訊清楚但不生硬，用卡面圖與重點摘要建立閱讀節奏。"),
])

heading(doc, "4. 分類頁概念與卡片規則", "分類頁承接首頁導流，重點是比較而不是陳列全部資訊，因此本版明確補上卡片高度與推薦理由規格。")
doc.add_picture(str(category_img), width=Inches(6.35))
kv_table(doc, [
    ("列表卡片內容", "卡面圖、卡名、銀行、主打回饋、優惠期間、推薦理由。"),
    ("卡片高度控制", "列表頁不允許內容無限展開；卡片高度需可控，避免少數長內容把後面卡片往下擠太多。"),
    ("推薦理由規則", "推薦理由在列表頁採二行摘要呈現，不做完整長文展開；超過二行需截斷，完整內容由詳情頁承接。"),
    ("列表頁定位", "列表頁的主要任務是快速比較，不是一次看完所有條件，因此細節限制與注意事項不在此頁大量展開。"),
])

heading(doc, "5. 詳情頁概念", "詳情頁負責補齊條件、期限與注意事項。它不是把列表資訊重貼，而是提供決策需要的完整脈絡。")
bullets(doc, [
    "卡片基本資訊：卡面、銀行、分類、狀態、標籤。",
    "主優惠資訊：主打內容、期限、門檻、回饋上限、注意事項。",
    "來源資訊：官方連結與備註，方便人工校對與後續爬蟲串接。",
])

heading(doc, "6. 後台概念與列表摘要設定", "MVP 的成敗在於後台能不能穩定維護資料。這一版補上列表摘要設定區塊，避免推薦理由完全依賴人工手寫。")
kv_table(doc, [
    ("管理模組", "儀表板、銀行管理、信用卡管理、分類管理、優惠管理、設定。"),
    ("摘要方案", "採方案 C：系統自動組摘要，後台可人工覆寫。"),
    ("顯示規則", "若 manual_summary 有值，前台優先顯示人工摘要；若為空，顯示系統自動組出的摘要。"),
])
grid_table(doc,
    ["欄位名稱", "輸入方式", "規格與用途"],
    [
        ("summary_mode", "單選", "可選「使用系統摘要」或「使用人工摘要」。"),
        ("target_audience", "下拉選單", "固定選項：日常消費、網購族、旅遊族、海外消費、報稅族、繳費族、新戶優先。"),
        ("highlight_1", "下拉選單", "固定選項：通路多、門檻低、回饋高、期限長、免登錄、海外適用、活動單純、回饋上限高。"),
        ("highlight_2", "下拉選單", "與 highlight_1 相同，用於組成第二個主優勢。"),
        ("manual_summary", "文字輸入（選填）", "供精選卡片人工覆寫摘要，保留編輯感。"),
        ("summary_preview", "唯讀預覽", "即時顯示前台卡片會看到的最終二行摘要。"),
    ],
    widths=[1.6, 1.4, 3.4],
)
kv_table(doc, [
    ("系統摘要句型", "先採簡單規則：適合{族群}，{優勢1}、{優勢2}；若欄位不足，再依實際有值欄位自動縮減句型。"),
    ("文案策略", "精選卡片可人工覆寫摘要；一般卡片直接使用系統摘要，降低維護成本。"),
])

doc.add_page_break()
heading(doc, "7. 資料結構規劃（補強版）", "資料模型需同時支援前台展示與後台維護，也要為未來爬蟲與多分類延伸保留空間。本版補上「欄位在後台如何被建立、編輯與使用」的規格。")
doc.add_picture(str(admin_img), width=Inches(6.35))
grid_table(doc,
    ["資料實體", "核心欄位", "後台對應與用途"],
    [
        ("banks", "name / slug / logo_url / website_url / is_active", "在銀行管理建立與維護，供信用卡資料關聯與前台銀行篩選使用。"),
        ("cards", "bank_id / name / slug / image_url / summary / is_active", "在信用卡管理建立卡片基本資訊，供列表頁卡名、銀行、卡面圖與詳情頁資料引用。"),
        ("categories", "name / slug / icon_name / sort_order / is_active", "在分類管理維護主導覽分類，用於首頁分類入口與分類列表頁路由。"),
        ("offers", "title / summary / description / start_date / end_date / reward_type / reward_value / reward_cap / min_spend / conditions / source_url / tags / is_featured / recommend_score / sort_order / is_published", "在優惠管理建立單筆優惠，並在單筆編輯畫面填寫摘要、排序、上架與詳情欄位。"),
        ("offer_cards", "offer_id / card_id", "在優惠編輯畫面以多選方式關聯信用卡，支援一個優惠掛多張卡。"),
    ],
    widths=[1.0, 2.6, 2.7],
)
heading(doc, "7.1 單筆優惠編輯畫面規格", "後台不只要能維護資料，還要讓欄位用途與前台影響一眼看懂。")
grid_table(doc,
    ["區塊", "欄位", "操作說明"],
    [
        ("基本資訊", "優惠標題 / 所屬分類 / 關聯信用卡", "決定前台標題、分類歸屬與卡片歸屬。"),
        ("列表摘要設定", "summary_mode / target_audience / highlight_1 / highlight_2 / manual_summary / summary_preview", "決定列表頁推薦理由如何產生與如何顯示。"),
        ("排序與上架設定", "is_featured / recommend_score / sort_order / is_published / updated_at", "決定前台排序優先順序與是否可見。"),
        ("詳情資料", "description / conditions / reward_cap / min_spend / source_url", "由詳情頁承接列表頁未展開的完整內容。"),
    ],
    widths=[1.2, 2.4, 2.7],
)

heading(doc, "8. 商業規則與排序（補強版）", "第一版不做黑盒推薦，而是用人工可控的方式建立可信任排序。本版補上排序欄位規格與後台操作方式。")
kv_table(doc, [
    ("前台顯示原則", "前台僅顯示 is_published 的資料；若超過 end_date 可標記為已過期。"),
    ("預設排序", "is_featured → recommend_score → sort_order → updated_at。"),
    ("排序目的", "先讓編輯能控制主推內容，再用分數與微調順序處理同類優惠的前後。"),
])
grid_table(doc,
    ["欄位", "資料型別", "後台輸入方式", "預設值", "用途說明", "排序影響 / 維護方式"],
    [
        ("is_featured", "Boolean", "切換開關", "false", "代表本站推薦。開啟後前台可顯示推薦標記。", "會影響排序；人工維護。"),
        ("recommend_score", "Integer", "數字輸入", "0", "推薦權重分數，分數越高排序越前。", "會影響排序；人工維護。"),
        ("sort_order", "Integer", "數字輸入", "0", "人工微調排序用，數值越小越前；當 recommend_score 相同時作為下一層排序。", "會影響排序；人工維護。"),
        ("updated_at", "Datetime", "系統自動帶入", "系統寫入", "紀錄最後更新時間，當其他條件相同時以最新更新時間排前面。", "會影響排序；系統自動更新。"),
    ],
    widths=[0.95, 0.9, 1.15, 0.7, 2.3, 1.0],
)
heading(doc, "8.1 排序操作規則", "排序欄位不應只出現在資料表說明裡，還要讓後台管理者知道實際怎麼用。")
bullets(doc, [
    "is_featured：用來決定是否為本站推薦；同類優惠中，先排在一般優惠前。",
    "recommend_score：用來評估該筆優惠是否值得優先露出；例如高分表示更適合首頁或分類頁前段。",
    "sort_order：當兩筆優惠推薦分數相同時，由管理者手動微調，數字越小越前。",
    "updated_at：當上述排序條件仍相同時，系統以最後更新時間作為最後排序依據。",
])

heading(doc, "9. 技術建議與開發順序", "以開發速度與後續擴充性平衡，建議前後端同專案推進。")
numbered(doc, [
    "技術棧：Next.js、Prisma、SQLite、Tailwind CSS；後續可平滑升級 PostgreSQL。",
    "先完成資料庫 schema 與後台登入，再做後台 CRUD，最後串接前台首頁、分類頁與詳情頁。",
    "完成測試資料後，進行 RWD 與閱讀節奏微調，確保桌機與手機都好讀。",
])

heading(doc, "10. 第二階段擴充", "第一版站穩後，再把效率工具與資料自動化能力加上來。")
bullets(doc, [
    "銀行官網爬蟲與半自動匯入。",
    "規則評分機制 + 人工微調推薦。",
    "SEO 強化、主題懶人包頁、更多分類與季節性專題。",
])

heading(doc, "11. 本版更新說明", "以下為本版（2026-06-03 / v3）相較上一版 credit-card-mvp-spec 與 v2 的新增與補強內容。")
bullets(doc, [
    "修正上一版出現的亂碼、問號替代字與非繁體中文內容。",
    "新增分類列表頁卡片高度控制規則，明確禁止內容無限展開。",
    "新增推薦理由二行摘要規格，並說明超過二行需截斷、完整內容由詳情頁承接。",
    "新增列表摘要設定區塊，補上 summary_mode、target_audience、highlight_1、highlight_2、manual_summary、summary_preview 欄位。",
    "補強第 7 節資料結構規劃，新增後台欄位如何建立、編輯與使用的說明。",
    "補強第 8 節商業規則與排序，新增 is_featured、recommend_score、sort_order、updated_at 的完整欄位規格。",
    "新增更詳細的後台單筆優惠編輯畫面概念，讓閱讀者可理解排序與摘要欄位的實際位置與用途。",
    "新版檔案另存為日期與版本號格式，未覆蓋原本的 credit-card-mvp-spec。",
])

doc.save(DOCX_PATH)
TEXT_DUMP.write_text("\n".join(p.text for p in doc.paragraphs if p.text.strip()), encoding="utf-8")
print(DOCX_PATH)
print(TEXT_DUMP)
print(category_img)
print(admin_img)
