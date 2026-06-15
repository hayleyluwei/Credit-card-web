# -*- coding: utf-8 -*-
from pathlib import Path
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

ROOT = Path(r"C:\Users\user\Documents\Credit card web project")
OUT_DIR = ROOT / "outputs" / "product-spec"
ASSET_DIR = OUT_DIR / "assets-v20260605-v6-zh"
DOCX_PATH = OUT_DIR / "credit-card-mvp-spec-v7-2026-06-05.docx"
TEXT_DUMP = OUT_DIR / "credit-card-mvp-spec-v7-2026-06-05-inspection.txt"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

PALETTE = {
    "ink": (33, 43, 54),
    "navy": (20, 45, 78),
    "teal": (46, 122, 142),
    "mist": (240, 245, 247),
    "line": (207, 215, 222),
    "gold": (201, 161, 77),
    "gray": (95, 104, 112),
    "white": (255, 255, 255),
}


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msjhbd.ttc" if bold else "C:/Windows/Fonts/msjh.ttc"),
        Path("C:/Windows/Fonts/mingliub.ttc" if bold else "C:/Windows/Fonts/mingliu.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


F_TITLE = font(42, True)
F_H1 = font(30, True)
F_H2 = font(22, True)
F_BODY = font(18, False)
F_SMALL = font(16, False)


def rounded(draw, box, fill, outline=None, radius=18, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def title_block(draw, title, subtitle):
    draw.text((60, 42), title, font=F_TITLE, fill=PALETTE["navy"])
    draw.text((60, 102), subtitle, font=F_BODY, fill=PALETTE["ink"])
    draw.line((60, 148, 1540, 148), fill=PALETTE["line"], width=3)


def create_category_rule_image(path: Path):
    img = Image.new("RGB", (1600, 980), PALETTE["white"])
    draw = ImageDraw.Draw(img)
    title_block(draw, "分類頁卡片規則：固定高度與二行摘要", "列表頁負責快速比較，因此推薦理由只顯示二行；超過內容由詳情頁承接，不在列表頁無限展開。")
    draw.text((60, 188), "現金回饋", font=F_H1, fill=PALETTE["navy"])
    draw.text((60, 228), "卡片高度需可控，避免少數長內容把後面卡片整片往下擠。", font=F_BODY, fill=PALETTE["ink"])
    rounded(draw, (60, 278, 1540, 352), PALETTE["mist"], PALETTE["line"])
    filters = ["銀行：全部", "僅看進行中", "排序：本站推薦優先", "摘要：固定二行"]
    x = 84
    for item in filters:
        rounded(draw, (x, 298, x + 260, 334), PALETTE["white"], radius=10)
        draw.text((x + 16, 306), item, font=F_SMALL, fill=PALETTE["ink"])
        x += 280
    rows = [
        ("1", "台新太陽卡", "台新銀行", "一般消費最高 5%", "推薦理由：適合日常消費，門檻低、規則清楚。", "若還有更多限制條件，改在詳情頁完整說明。", True),
        ("2", "國泰 CUBE 卡", "國泰世華", "指定通路最高 3%", "推薦理由：適合數位消費，主題切換邏輯明確。", "列表頁只保留比較所需重點，不延伸長文。", False),
        ("3", "富邦 J 卡", "台北富邦", "旅遊場景最高 5%", "推薦理由：旅日場景鮮明，適合作為主題整理卡片。", "完整活動細節與限制條件交由詳情頁承接。", False),
    ]
    y = 388
    for rank, title, bank, reward, line1, line2, featured in rows:
        rounded(draw, (60, y, 1540, y + 148), PALETTE["white"], PALETTE["line"], radius=20)
        rounded(draw, (82, y + 28, 138, y + 84), PALETTE["navy"], radius=16)
        draw.text((104, y + 35), rank, font=F_H2, fill=PALETTE["white"])
        rounded(draw, (166, y + 20, 314, y + 118), (229, 236, 241), radius=16)
        draw.text((208, y + 56), "卡面圖", font=F_H2, fill=PALETTE["navy"])
        draw.text((344, y + 22), title, font=F_H2, fill=PALETTE["ink"])
        draw.text((344, y + 52), bank, font=F_SMALL, fill=PALETTE["gray"])
        draw.text((344, y + 82), line1, font=F_SMALL, fill=PALETTE["gray"])
        draw.text((344, y + 108), line2, font=F_SMALL, fill=PALETTE["gray"])
        draw.text((1060, y + 28), reward, font=F_H1, fill=PALETTE["teal"])
        draw.text((1060, y + 74), "優惠期間：2026/06/01 - 2026/12/31", font=F_SMALL, fill=PALETTE["ink"])
        if featured:
            rounded(draw, (1328, y + 28, 1498, y + 64), PALETTE["gold"], radius=16)
            draw.text((1354, y + 36), "本站推薦", font=F_SMALL, fill=PALETTE["ink"])
        y += 170
    img.save(path)


def create_detail_page_image(path: Path):
    img = Image.new("RGB", (1600, 940), PALETTE["white"])
    draw = ImageDraw.Draw(img)
    title_block(draw, "詳情頁示意圖：完整條件、期間與注意事項", "詳情頁承接列表頁未展開的資訊，讓使用者能在決策前看到完整優惠條件與官方來源。")
    rounded(draw, (60, 190, 420, 840), PALETTE["white"], PALETTE["line"], radius=22)
    rounded(draw, (98, 232, 382, 428), (229, 236, 241), radius=18)
    draw.text((184, 314), "卡面圖", font=F_H1, fill=PALETTE["navy"])
    draw.text((98, 478), "國泰 CUBE 卡", font=F_H1, fill=PALETTE["ink"])
    draw.text((98, 522), "銀行：國泰世華", font=F_BODY, fill=PALETTE["gray"])
    draw.text((98, 560), "分類：現金回饋", font=F_BODY, fill=PALETTE["gray"])
    draw.text((98, 598), "狀態：進行中", font=F_BODY, fill=PALETTE["teal"])
    rounded(draw, (98, 646, 198, 686), PALETTE["navy"], radius=14)
    draw.text((120, 654), "指定通路", font=F_SMALL, fill=PALETTE["white"])
    rounded(draw, (214, 646, 330, 686), PALETTE["teal"], radius=14)
    draw.text((238, 654), "數位消費", font=F_SMALL, fill=PALETTE["white"])

    rounded(draw, (450, 190, 1540, 840), PALETTE["white"], PALETTE["line"], radius=22)
    draw.text((490, 228), "主打優惠：指定通路最高 3% 回饋", font=F_H1, fill=PALETTE["navy"])
    sections = [
        ("優惠期間", "2026/06/01 - 2026/12/31"),
        ("適用條件", "需於 App 切換對應權益方案，並符合指定通路消費條件。"),
        ("回饋上限", "每月回饋上限 300 點，超過上限不再累積。"),
        ("注意事項", "若活動條件或通路名單變更，以官方公告為準，詳情頁需保留完整說明。"),
        ("官方連結", "https://example-bank.com/cube-offer"),
    ]
    y = 296
    for label, value in sections:
        rounded(draw, (490, y, 1502, y + 92), PALETTE["mist"], radius=18)
        draw.text((520, y + 18), label, font=F_H2, fill=PALETTE["ink"])
        draw.text((778, y + 24), value, font=F_SMALL, fill=PALETTE["gray"])
        y += 108
    img.save(path)


def create_admin_editor_image(path: Path):
    img = Image.new("RGB", (1600, 1260), PALETTE["white"])
    draw = ImageDraw.Draw(img)
    title_block(draw, "後台畫面補強：單筆優惠編輯與排序設定", "這張示意圖必須和文件正文對得上，因此欄位名稱、輸入方式、畫面值與用途說明都要逐項對應。")
    rounded(draw, (60, 184, 300, 1200), PALETTE["navy"], radius=26)
    menu = ["儀表板", "銀行管理", "信用卡管理", "分類管理", "優惠管理", "設定"]
    y = 238
    for item in menu:
        rounded(draw, (82, y, 278, y + 48), PALETTE["white"], radius=14)
        draw.text((110, y + 11), item, font=F_BODY, fill=PALETTE["navy"])
        y += 74
    rounded(draw, (330, 184, 1540, 1200), PALETTE["white"], PALETTE["line"], radius=24)
    draw.text((364, 220), "優惠編輯：CUBE 指定通路回饋", font=F_H1, fill=PALETTE["navy"])

    sections = [
        (
            "基本資訊",
            [
                ("優惠標題", "文字輸入", "CUBE 指定通路回饋", "控制前台卡片標題與詳情頁主標"),
                ("所屬分類", "下拉選單", "現金回饋", "決定分類頁歸屬"),
                ("關聯信用卡", "多選欄位", "CUBE 卡、CUBE COMBO 卡", "支援一個優惠掛多張卡"),
            ],
        ),
        (
            "列表摘要設定",
            [
                ("summary_mode", "單選按鈕", "使用系統摘要", "決定使用系統摘要或人工摘要"),
                ("target_audience", "下拉選單", "數位消費", "用於組成推薦理由中的適合族群"),
                ("highlight_1", "下拉選單", "通路多", "用於組成推薦理由中的第一個主優勢"),
                ("highlight_2", "下拉選單", "活動單純", "用於組成推薦理由中的第二個主優勢"),
                ("manual_summary", "文字輸入（選填）", "若空白則不覆寫", "若有值，前台優先顯示人工摘要"),
                ("summary_preview", "唯讀預覽", "適合數位消費，通路多、活動單純", "預覽前台列表卡片實際看到的二行摘要"),
            ],
        ),
        (
            "排序與上架設定",
            [
                ("is_featured", "開關", "開啟", "控制本站推薦標記，並優先排序"),
                ("recommend_score", "數字輸入 0-100", "95", "分數越高越前面"),
                ("sort_order", "整數輸入", "10", "同分時人工微調，數字越小越前"),
                ("is_published", "開關", "開啟", "控制是否可出現在前台"),
                ("updated_at", "系統自動帶入", "2026/06/04 09:30", "其他排序條件相同時，作為最後排序依據"),
            ],
        ),
    ]

    y = 280
    for sec_title, rows in sections:
        height = 100 + len(rows) * 52
        rounded(draw, (364, y, 1508, y + height), PALETTE["mist"], PALETTE["line"], radius=18)
        draw.text((390, y + 18), sec_title, font=F_H2, fill=PALETTE["navy"])
        ry = y + 58
        draw.text((398, ry + 8), "欄位名稱", font=F_SMALL, fill=PALETTE["gray"])
        draw.text((620, ry + 8), "輸入方式", font=F_SMALL, fill=PALETTE["gray"])
        draw.text((860, ry + 8), "畫面值 / 預設值", font=F_SMALL, fill=PALETTE["gray"])
        draw.text((1164, ry + 8), "用途說明", font=F_SMALL, fill=PALETTE["gray"])
        ry += 36
        draw.line((390, ry, 1480, ry), fill=PALETTE["line"], width=1)
        for label, input_type, sample_value, desc in rows:
            draw.text((398, ry + 10), label, font=F_SMALL, fill=PALETTE["ink"])
            draw.text((620, ry + 10), input_type, font=F_SMALL, fill=PALETTE["gray"])
            draw.text((860, ry + 10), sample_value, font=F_SMALL, fill=PALETTE["ink"])
            draw.text((1164, ry + 10), desc, font=F_SMALL, fill=PALETTE["gray"])
            ry += 52
            draw.line((390, ry, 1480, ry), fill=PALETTE["line"], width=1)
        y = ry + 22
    img.save(path)


def create_admin_module_image(path: Path, active_menu: str, title: str, subtitle: str, columns, rows, footer_note: str):
    img = Image.new("RGB", (1600, 1040), PALETTE["white"])
    draw = ImageDraw.Draw(img)
    title_block(draw, title, subtitle)
    rounded(draw, (60, 184, 300, 980), PALETTE["navy"], radius=26)
    menu = ["儀表板", "銀行管理", "信用卡管理", "分類管理", "優惠管理", "設定"]
    y = 238
    for item in menu:
        fill = PALETTE["gold"] if item == active_menu else PALETTE["white"]
        text_color = PALETTE["ink"] if item == active_menu else PALETTE["navy"]
        rounded(draw, (82, y, 278, y + 48), fill, radius=14)
        draw.text((110, y + 11), item, font=F_BODY, fill=text_color)
        y += 74

    rounded(draw, (330, 184, 1540, 980), PALETTE["white"], PALETTE["line"], radius=24)
    draw.text((364, 220), title, font=F_H1, fill=PALETTE["navy"])
    rounded(draw, (1210, 218, 1490, 268), PALETTE["teal"], radius=16)
    draw.text((1260, 232), "新增 / 儲存", font=F_BODY, fill=PALETTE["white"])

    table_left = 364
    table_top = 304
    table_right = 1508
    row_h = 66
    col_w = (table_right - table_left) // len(columns)
    rounded(draw, (table_left, table_top, table_right, table_top + row_h), PALETTE["mist"], PALETTE["line"], radius=14)
    for idx, col in enumerate(columns):
        x = table_left + idx * col_w + 14
        draw.text((x, table_top + 22), col, font=F_SMALL, fill=PALETTE["navy"])

    y = table_top + row_h
    for row in rows:
        fill = PALETTE["white"] if (y // row_h) % 2 == 0 else (247, 250, 251)
        draw.rectangle((table_left, y, table_right, y + row_h), fill=fill)
        draw.line((table_left, y + row_h, table_right, y + row_h), fill=PALETTE["line"], width=1)
        for idx, value in enumerate(row):
            x = table_left + idx * col_w + 14
            draw.text((x, y + 20), value, font=F_SMALL, fill=PALETTE["ink"])
        y += row_h

    rounded(draw, (364, 888, 1508, 948), PALETTE["mist"], PALETTE["line"], radius=16)
    draw.text((390, 908), footer_note, font=F_BODY, fill=PALETTE["gray"])
    img.save(path)


def create_bank_management_image(path: Path):
    create_admin_module_image(
        path,
        "銀行管理",
        "銀行管理：銀行資料列表",
        "點擊左側銀行管理後，管理者可維護銀行基本資料，供信用卡關聯與前台銀行篩選使用。",
        ["銀行名稱", "slug", "Logo", "官網", "狀態", "更新時間", "操作"],
        [
            ["台新銀行", "taishin", "已上傳", "taishinbank.com.tw", "啟用", "2026/06/05", "編輯 / 停用"],
            ["國泰世華", "cathay", "已上傳", "cathaybk.com.tw", "啟用", "2026/06/04", "編輯 / 停用"],
            ["台北富邦", "fubon", "待補", "fubon.com", "停用", "2026/06/01", "編輯 / 刪除"],
        ],
        "銀行資料會被信用卡資料引用，也會用於前台銀行篩選。",
    )


def create_card_management_image(path: Path):
    create_admin_module_image(
        path,
        "信用卡管理",
        "信用卡管理：卡片資料列表",
        "點擊左側信用卡管理後，管理者維護卡本身資料，不直接編輯單筆優惠內容。",
        ["卡名", "所屬銀行", "卡面圖", "卡片摘要", "狀態", "關聯優惠", "操作"],
        [
            ["CUBE 卡", "國泰世華", "已上傳", "數位通路彈性切換", "啟用", "4", "編輯 / 停用"],
            ["太陽卡", "台新銀行", "已上傳", "日常消費回饋", "啟用", "3", "編輯 / 停用"],
            ["J 卡", "台北富邦", "已上傳", "旅遊與海外場景", "啟用", "2", "編輯 / 停用"],
        ],
        "信用卡管理維護卡名、銀行、卡面圖與摘要；優惠內容由優惠管理維護。",
    )


def create_category_management_image(path: Path):
    create_admin_module_image(
        path,
        "分類管理",
        "分類管理：前台分類列表",
        "點擊左側分類管理後，管理者維護前台分類入口；分類主要掛在優惠上，不直接掛在信用卡上。",
        ["分類名稱", "slug", "icon", "排序", "狀態", "前台顯示", "操作"],
        [
            ["現金回饋", "cashback", "cash", "10", "啟用", "顯示", "編輯 / 停用"],
            ["網購", "online", "cart", "20", "啟用", "顯示", "編輯 / 停用"],
            ["旅遊", "travel", "plane", "30", "啟用", "顯示", "編輯 / 停用"],
            ["新戶首刷", "new-user", "gift", "40", "啟用", "顯示", "編輯 / 停用"],
        ],
        "一張信用卡可透過不同優惠出現在多個分類頁。",
    )


def create_offer_management_image(path: Path):
    create_admin_module_image(
        path,
        "優惠管理",
        "優惠管理：優惠列表",
        "點擊左側優惠管理後，管理者先在列表查找優惠，再進入單筆優惠編輯畫面維護內容。",
        ["優惠標題", "分類", "關聯信用卡", "期間", "上架", "推薦 / 分數", "操作"],
        [
            ["CUBE 指定通路回饋", "現金回饋", "CUBE 卡", "06/01-12/31", "上架", "是 / 95", "編輯 / 預覽"],
            ["太陽卡日常消費", "日常消費", "太陽卡", "06/01-09/30", "草稿", "否 / 70", "編輯 / 下架"],
            ["J 卡旅遊回饋", "旅遊", "J 卡", "07/01-12/31", "上架", "否 / 82", "編輯 / 預覽"],
        ],
        "優惠列表負責查找與管理；單筆編輯頁負責基本資訊、摘要、排序、上架與詳情資料。",
    )


def create_settings_image(path: Path):
    create_admin_module_image(
        path,
        "設定",
        "設定：全站參數",
        "點擊左側設定後，管理者維護全站參數；此頁不是內容資料管理頁。",
        ["設定項目", "欄位值", "輸入方式", "用途"],
        [
            ["網站名稱", "信用卡優惠查詢站", "文字輸入", "前台與 SEO 預設名稱"],
            ["SEO 預設標題", "信用卡優惠比較", "文字輸入", "未自訂頁面的預設標題"],
            ["SEO 預設描述", "整理最新信用卡優惠", "文字輸入", "未自訂頁面的預設描述"],
            ["首頁精選數量", "6", "數字輸入", "控制首頁精選卡片數"],
            ["每頁顯示筆數", "20", "數字輸入", "控制分類列表分頁"],
            ["顯示過期優惠", "否", "開關", "控制前台是否保留過期優惠"],
            ["管理員帳號", "admin@example.com", "帳號設定", "後台登入與權限管理"],
        ],
        "MVP 初期設定頁保持簡單，只維護必要的全站參數。",
    )


def set_fill(cell, color_hex: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def style(run, size=10.5, bold=False, color="212B36"):
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def heading(doc, text, intro=None):
    p = doc.add_paragraph()
    style(p.add_run(text), 18, True, "14334E")
    if intro:
        p2 = doc.add_paragraph()
        style(p2.add_run(intro), 10.5, False, "394651")


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        style(p.add_run(item), 10.5)


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        style(p.add_run(item), 10.5)


def kv_table(doc, rows, col1=2.0, col2=4.3):
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(col1)
    table.columns[1].width = Inches(col2)
    for left, right in rows:
        cells = table.add_row().cells
        set_fill(cells[0], "F0F5F7")
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        style(cells[0].paragraphs[0].add_run(left), 10, True, "14334E")
        style(cells[1].paragraphs[0].add_run(right), 10, False, "212B36")
    doc.add_paragraph()


def grid_table(doc, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.autofit = False
    if widths:
        for idx, width in enumerate(widths):
            table.columns[idx].width = Inches(width)
    for idx, header in enumerate(headers):
        set_fill(table.cell(0, idx), "DDE7ED")
        style(table.cell(0, idx).paragraphs[0].add_run(header), 9.5, True, "14334E")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            style(cells[idx].paragraphs[0].add_run(value), 9.3, False, "212B36")
    doc.add_paragraph()


category_img = ASSET_DIR / "分類頁卡片規則-2026-06-05-v6.png"
detail_img = ASSET_DIR / "詳情頁示意圖-2026-06-05-v6.png"
admin_img = ASSET_DIR / "後台單筆優惠編輯畫面-2026-06-05-v6.png"
bank_admin_img = ASSET_DIR / "銀行管理畫面-2026-06-05-v6.png"
card_admin_img = ASSET_DIR / "信用卡管理畫面-2026-06-05-v6.png"
category_admin_img = ASSET_DIR / "分類管理畫面-2026-06-05-v6.png"
offer_admin_img = ASSET_DIR / "優惠管理畫面-2026-06-05-v6.png"
settings_admin_img = ASSET_DIR / "設定畫面-2026-06-05-v6.png"
frontend_offer_img = ASSET_DIR / "frontend-offer-display-2026-06-05.png"
create_category_rule_image(category_img)
create_detail_page_image(detail_img)
create_admin_editor_image(admin_img)
create_bank_management_image(bank_admin_img)
create_card_management_image(card_admin_img)
create_category_management_image(category_admin_img)
create_offer_management_image(offer_admin_img)
create_settings_image(settings_admin_img)

doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.75)
section.bottom_margin = Inches(0.7)
section.left_margin = Inches(0.75)
section.right_margin = Inches(0.75)
normal = doc.styles["Normal"]
normal.font.name = "Microsoft JhengHei"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
normal.font.size = Pt(10.5)

p = doc.add_paragraph()
style(p.add_run("信用卡優惠查詢網站 MVP 規格書"), 24, True, "14334E")
p2 = doc.add_paragraph()
style(p2.add_run("版本：2026-06-05 / v7\n本版在 v6 基礎上補入後台優惠資料對應前台分類頁與詳情頁的畫面圖，讓後台欄位如何落到前台展示更加直觀。"), 11, False, "394651")

heading(doc, "1. 產品目標", "這個網站不是冷資料庫，而是讓使用者可以快速理解哪張卡值得看、為什麼值得看的整理型入口站。")
bullets(doc, [
    "前台提供公開查詢，使用者可從分類、關鍵字、銀行與是否進行中快速篩選。",
    "首頁偏導覽與整理感，透過圖文卡片、分類入口與精選區塊降低閱讀負擔。",
    "後台優先解決資料治理，讓管理者能維護銀行、信用卡、優惠、分類與推薦排序。",
])

heading(doc, "2. MVP 範圍", "第一版聚焦在能上線、能維護、能擴充，而不是一次做到最完整。")
numbered(doc, [
    "前台：首頁、分類列表頁、優惠詳情頁、搜尋與篩選。",
    "後台：單一管理員登入，銀行 / 卡片 / 分類 / 優惠資料 CRUD。",
    "排序：採手動推薦權重，後台可編輯 is_featured、recommend_score、sort_order。",
    "資料：先人工建檔，預留未來接銀行官網爬蟲與規則評分。",
])

heading(doc, "3. 首頁概念", "首頁是懶人包式入口頁。左側讓使用者快速切入優惠主題，右側集中展示最新與值得先看的信用卡。")
kv_table(doc, [
    ("核心任務", "讓使用者 10 秒內找到自己想看的優惠主題。"),
    ("主要元件", "分類導覽、搜尋框、精選卡片、進行中與推薦標籤。"),
    ("視覺方向", "資訊清楚但不生硬，用卡面圖與重點摘要建立閱讀節奏。"),
])

heading(doc, "4. 分類頁概念與卡片規則", "分類頁承接首頁導流，重點是比較而不是陳列全部資訊，因此本版明確補上卡片高度與推薦理由規格。")
doc.add_picture(str(category_img), width=Inches(6.35))
kv_table(doc, [
    ("列表卡片內容", "卡面圖、卡名、銀行、主打回饋、優惠期間、推薦理由。"),
    ("卡片高度控制", "列表頁不允許內容無限展開；卡片高度需可控，避免少數長內容把後面卡片往下擠太多。"),
    ("推薦理由規則", "推薦理由在列表頁採二行摘要呈現，不做完整長文展開；超過二行需截斷，完整內容由詳情頁承接。"),
    ("列表頁定位", "列表頁的主要任務是快速比較，不是一次看完所有條件，因此細節限制與注意事項不在此頁大量展開。"),
])

heading(doc, "5. 詳情頁概念", "詳情頁負責補齊條件、期限與注意事項。它不是把列表資訊重貼，而是提供決策需要的完整脈絡。")
doc.add_picture(str(detail_img), width=Inches(6.35))
bullets(doc, [
    "卡片基本資訊：卡面、銀行、分類、狀態、標籤。",
    "主優惠資訊：主打內容、期限、適用條件、回饋上限、注意事項。",
    "來源資訊：官方連結與備註，方便人工校對與後續爬蟲串接。",
])

heading(doc, "6. 後台概念與列表摘要設定", "MVP 的成敗在於後台能不能穩定維護資料。這一版補上列表摘要設定區塊，避免推薦理由完全依賴人工手寫。")
kv_table(doc, [
    ("管理模組", "儀表板、銀行管理、信用卡管理、分類管理、優惠管理、設定。"),
    ("摘要方案", "採方案 C：系統自動組摘要，後台可人工覆寫。"),
    ("顯示規則", "若 manual_summary 有值，前台優先顯示人工摘要；若為空，顯示系統自動組出的摘要。"),
])
grid_table(doc,
    ["欄位名稱", "輸入方式", "規格與用途"],
    [
        ("summary_mode", "單選", "可選「使用系統摘要」或「使用人工摘要」。"),
        ("target_audience", "下拉選單", "固定選項：日常消費、網購族、旅遊族、海外消費、報稅族、繳費族、新戶優先。"),
        ("highlight_1", "下拉選單", "固定選項：通路多、門檻低、回饋高、期限長、免登錄、海外適用、活動單純、回饋上限高。"),
        ("highlight_2", "下拉選單", "與 highlight_1 使用相同選項，用於組成第二個主優勢。"),
        ("manual_summary", "文字輸入（選填）", "供精選卡片人工覆寫摘要，保留編輯感。"),
        ("summary_preview", "唯讀預覽", "即時顯示前台卡片會看到的最終二行摘要。"),
    ],
    widths=[1.6, 1.4, 3.4],
)
kv_table(doc, [
    ("系統摘要句型", "先採簡單規則：適合{族群}，{優勢1}、{優勢2}；若欄位不足，再依實際有值欄位自動縮減句型。"),
    ("文案策略", "精選卡片可人工覆寫摘要；一般卡片直接使用系統摘要，降低維護成本。"),
])

doc.add_page_break()
heading(doc, "7. 資料結構規劃（補強版）", "資料模型需同時支援前台展示與後台維護，也要為未來爬蟲與多分類延伸保留空間。本版補上「欄位在後台如何被建立、編輯與使用」的規格。")
doc.add_picture(str(admin_img), width=Inches(6.35))
grid_table(doc,
    ["資料實體", "核心欄位", "後台對應與用途"],
    [
        ("banks", "name / slug / logo_url / website_url / is_active", "在銀行管理建立與維護，供信用卡資料關聯與前台銀行篩選使用。"),
        ("cards", "bank_id / name / slug / image_url / summary / is_active", "在信用卡管理建立卡片基本資訊，供列表頁卡名、銀行、卡面圖與詳情頁資料引用。"),
        ("categories", "name / slug / icon_name / sort_order / is_active", "在分類管理維護主導覽分類，用於首頁分類入口與分類列表頁路由。"),
        ("offers", "title / summary / summary_mode / target_audience / highlight_1 / highlight_2 / manual_summary / summary_preview / description / start_date / end_date / reward_type / reward_value / reward_cap / min_spend / conditions / source_url / tags / is_featured / recommend_score / sort_order / is_published", "在優惠管理建立單筆優惠，並在單筆編輯畫面填寫摘要、排序、上架與詳情欄位；摘要相關欄位用於產生前台列表頁推薦理由。"),
        ("offer_cards", "offer_id / card_id", "優惠與信用卡的多對多中間表，在優惠編輯畫面以多選方式關聯信用卡，支援一個優惠掛多張卡，也支援一張卡關聯多個優惠。"),
    ],
    widths=[1.0, 2.6, 2.7],
)
kv_table(doc, [
    ("摘要欄位用途", "summary 為系統或後台可保留的列表摘要資料；manual_summary 是人工覆寫文字；summary_preview 是依 summary_mode、target_audience、highlight_1、highlight_2 與 manual_summary 計算出的前台二行摘要預覽。"),
    ("分類關聯規則", "分類歸屬以 offers 為主，不直接掛在 cards；一張信用卡可以透過不同優惠出現在現金回饋、網購、旅遊等多個分類中。"),
    ("offer_cards 用途", "offer_cards 負責優惠與信用卡的多對多關聯，避免開發者誤以為一張信用卡只能屬於單一分類。"),
])
heading(doc, "7.1 單筆優惠編輯畫面規格", "後台不只要能維護資料，還要讓欄位用途與前台影響一眼看懂。")
grid_table(doc,
    ["區塊", "欄位", "操作說明"],
    [
        ("基本資訊", "優惠標題 / 所屬分類 / 關聯信用卡", "決定前台標題、分類歸屬與卡片歸屬。"),
        ("列表摘要設定", "summary_mode / target_audience / highlight_1 / highlight_2 / manual_summary / summary_preview", "決定列表頁推薦理由如何產生與如何顯示。"),
        ("排序與上架設定", "is_featured / recommend_score / sort_order / is_published / updated_at", "決定前台排序優先順序與是否可見。"),
        ("詳情資料", "description / conditions / reward_cap / min_spend / source_url", "由詳情頁承接列表頁未展開的完整內容。"),
    ],
    widths=[1.2, 2.4, 2.7],
)

heading(doc, "7.2 後台左側選單頁面規格", "左側選單不只是示意導覽，需明確定義每個模組點擊後的頁面內容，讓開發者能對照實作後台 CRUD 與設定畫面。")
grid_table(doc,
    ["選單", "點擊後畫面", "主要內容"],
    [
        ("儀表板", "後台首頁總覽", "顯示目前上架優惠數、草稿優惠數、已過期優惠數、信用卡總數、銀行總數、即將到期或待檢查優惠、最近更新優惠清單，並提供新增優惠、新增信用卡、查看前台等快捷操作。"),
        ("銀行管理", "銀行資料列表與銀行表單", "列表包含銀行名稱、銀行 slug、Logo、官網連結、啟用狀態、建立或更新時間，以及新增、編輯、停用、刪除操作；表單可維護銀行名稱、Logo、官網與啟用狀態。"),
        ("信用卡管理", "信用卡資料列表與信用卡表單", "列表包含卡名、所屬銀行、卡面圖、卡片摘要、啟用狀態、關聯優惠數，以及新增、編輯、停用、刪除操作；此模組維護卡本身，不維護單筆優惠內容。"),
        ("分類管理", "分類資料列表與分類表單", "列表包含分類名稱、slug、icon、排序、啟用狀態、前台是否顯示，以及新增、編輯、停用、刪除操作；分類可包含現金回饋、網購、旅遊、海外消費、繳費、新戶首刷等。"),
        ("優惠管理", "優惠列表與單筆優惠編輯畫面", "列表包含優惠標題、所屬分類、關聯信用卡、優惠期間、上架狀態、是否本站推薦、推薦分數、人工排序、更新時間，以及新增、編輯、預覽、下架、刪除操作；單筆編輯畫面包含基本資訊、列表摘要設定、排序與上架設定、詳情資料。"),
        ("設定", "全站設定頁", "維護網站名稱、SEO 預設標題、SEO 預設描述、首頁精選數量、分類頁每頁顯示筆數、是否顯示過期優惠、管理員帳號設定與基本系統參數；初期保持簡單，不擴充過多進階功能。"),
    ],
    widths=[0.9, 1.5, 3.9],
)

heading(doc, "7.3 後台各模組畫面圖", "以下畫面圖對應左側選單點擊後的主要頁面，補充各模組列表、欄位與操作方式；既有單筆優惠編輯示意圖仍作為優惠編輯頁的細部畫面。")
doc.add_picture(str(bank_admin_img), width=Inches(6.35))
doc.add_paragraph()
doc.add_picture(str(card_admin_img), width=Inches(6.35))
doc.add_paragraph()
doc.add_picture(str(category_admin_img), width=Inches(6.35))
doc.add_page_break()
doc.add_picture(str(offer_admin_img), width=Inches(6.35))
doc.add_paragraph()
doc.add_picture(str(settings_admin_img), width=Inches(6.35))

heading(doc, "7.4 後台優惠資料與前台顯示對應圖", "以下畫面圖說明後台優惠列表中的優惠標題、分類、關聯信用卡、期間、上架狀態與推薦分數，如何轉成前台分類頁卡片與詳情頁內容。")
doc.add_picture(str(frontend_offer_img), width=Inches(6.35))
kv_table(doc, [
    ("前台分類頁", "使用者在對應分類頁會看到已上架優惠的卡片摘要，顯示卡名、銀行、分類、優惠期間、二行推薦理由與主打回饋。"),
    ("前台詳情頁", "使用者點擊列表卡片後，會看到完整主標、卡面圖、期間、適用條件、回饋上限、注意事項與官方來源。"),
    ("後台對前台的關聯", "後台的優惠標題、分類、關聯信用卡、期間、上架狀態與推薦 / 分數，會共同影響前台出現的位置、排序與顯示內容。"),
])

heading(doc, "8. 商業規則與排序（補強版）", "第一版不做黑盒推薦，而是用人工可控的方式建立可信任排序。本版補上排序欄位規格與後台操作方式。")
kv_table(doc, [
    ("前台顯示原則", "前台僅顯示 is_published 的資料；若超過 end_date 可標記為已過期。"),
    ("預設排序", "is_featured → recommend_score → sort_order → updated_at。"),
    ("排序目的", "先讓編輯能控制主推內容，再用分數與微調順序處理同類優惠的前後。"),
])
grid_table(doc,
    ["欄位", "資料型別", "後台輸入方式", "預設值", "用途說明", "排序影響 / 維護方式"],
    [
        ("is_featured", "Boolean", "切換開關", "false", "代表本站推薦。開啟後前台可顯示推薦標記。", "會影響排序；人工維護。"),
        ("recommend_score", "Integer", "數字輸入", "0", "推薦權重分數，分數越高排序越前。", "會影響排序；人工維護。"),
        ("sort_order", "Integer", "數字輸入", "0", "人工微調排序用，數值越小越前；當 recommend_score 相同時作為下一層排序。", "會影響排序；人工維護。"),
        ("updated_at", "Datetime", "系統自動帶入", "系統寫入", "紀錄最後更新時間，當其他條件相同時以最新更新時間排前面。", "會影響排序；系統自動更新。"),
    ],
    widths=[0.95, 0.9, 1.15, 0.7, 2.3, 1.0],
)
heading(doc, "8.1 排序操作規則", "排序欄位不應只出現在資料表說明裡，還要讓後台管理者知道實際怎麼用。")
bullets(doc, [
    "is_featured：用來決定是否為本站推薦；同類優惠中，先排在一般優惠前。",
    "recommend_score：用來評估該筆優惠是否值得優先露出；例如高分表示更適合首頁或分類頁前段。",
    "sort_order：當兩筆優惠推薦分數相同時，由管理者手動微調，數字越小越前。",
    "updated_at：當上述排序條件仍相同時，系統以最後更新時間作為最後排序依據。",
])

heading(doc, "9. 技術建議與開發順序", "以開發速度與後續擴充性平衡，建議前後端同專案推進。")
numbered(doc, [
    "技術棧：Next.js、Prisma、SQLite、Tailwind CSS；後續可平滑升級 PostgreSQL。",
    "先完成資料庫 schema 與後台登入，再做後台 CRUD，最後串接前台首頁、分類頁與詳情頁。",
    "完成測試資料後，進行 RWD 與閱讀節奏微調，確保桌機與手機都好讀。",
])

heading(doc, "9.1 SEO 與 AI 搜尋可見性規格", "網站規劃需讓搜尋引擎與 AI 搜尋系統容易理解、收錄與引用，因此重要內容必須以可讀 HTML 文字、固定 URL 與結構化資料呈現。")
grid_table(doc,
    ["項目", "規格"],
    [
        ("固定 URL", "首頁、分類頁、銀行頁、信用卡頁與優惠詳情頁都需有固定可收錄 URL，不只依賴前端篩選狀態。"),
        ("SEO metadata", "分類頁、銀行頁、信用卡頁與優惠詳情頁需可設定或產生 SEO title、SEO description、slug 與 canonical URL。"),
        ("sitemap / robots", "需建立 sitemap.xml，收錄首頁、分類頁、銀行頁、信用卡頁與優惠詳情頁；需建立 robots.txt，避免誤擋重要頁面。"),
        ("詳情頁文字內容", "優惠詳情頁需以清楚段落呈現適合族群、主打回饋、優惠期間、適用條件、回饋上限、注意事項、官方來源、最後更新時間與是否人工校對。"),
        ("結構化資料", "初期規劃 JSON-LD：WebSite、Organization、BreadcrumbList、WebPage 或 Article；若頁面有常見問題區塊，加入 FAQPage。結構化資料必須和頁面可見內容一致。"),
        ("FAQ 區塊", "分類頁與重要優惠頁可加入 FAQ，例如現金回饋信用卡怎麼選、回饋上限是什麼、新戶禮和一般回饋是否可同時取得。"),
        ("圖片 alt text", "卡面圖、銀行 Logo 與示意圖需具備有意義的 alt text，避免重要資訊只存在圖片中。"),
        ("過期優惠", "過期優惠可保留頁面，但需清楚標示已過期，避免使用者與搜尋系統誤解。"),
        ("後台欄位", "分類管理預留 SEO title、SEO description、分類說明、FAQ；銀行管理預留銀行介紹、官網、Logo alt；信用卡管理預留卡片介紹、適合族群、卡面 alt；優惠管理預留來源 URL、最後校對時間、SEO title、SEO description、FAQ。"),
    ],
    widths=[1.35, 4.95],
)

heading(doc, "10. 第二階段擴充", "第一版站穩後，再把效率工具與資料自動化能力加上來。")
bullets(doc, [
    "銀行官網爬蟲與半自動匯入。",
    "規則評分機制 + 人工微調推薦。",
    "SEO 強化、主題懶人包頁、更多分類與季節性專題。",
])

heading(doc, "11. 本版更新說明", "以下為本版（2026-06-05 / v7）相較上一版 credit-card-mvp-spec-v6-2026-06-05 的新增、修正與補強內容。")
kv_table(doc, [
    ("新增", "新增 7.4 後台優惠資料與前台顯示對應圖，說明優惠管理列表中的一筆資料如何轉成前台分類頁卡片與優惠詳情頁。"),
    ("修正", "本版未修改既有技術棧、SEO / AI 搜尋規格與後台模組畫面，只補入前台對應圖。"),
    ("補強", "補強後台優惠列表、前台分類頁與前台詳情頁之間的資料對應關係，讓閱讀者可直接理解後台欄位如何影響前台展示。"),
    ("版本處理", "新版檔案另存為 credit-card-mvp-spec-v7-2026-06-05.docx，未覆蓋原本的 v6。"),
])

doc.save(DOCX_PATH)
TEXT_DUMP.write_text("\n".join(p.text for p in doc.paragraphs if p.text.strip()), encoding="utf-8")
print(DOCX_PATH)
print(TEXT_DUMP)
print(category_img)
print(detail_img)
print(admin_img)
print(bank_admin_img)
print(card_admin_img)
print(category_admin_img)
print(offer_admin_img)
print(settings_admin_img)
print(frontend_offer_img)
