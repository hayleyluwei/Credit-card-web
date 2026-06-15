# -*- coding: utf-8 -*-
from pathlib import Path
from shutil import copyfile

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(r"C:\Users\user\Documents\Credit card web project")
OUT_DIR = ROOT / "outputs" / "product-spec"
SRC_DOCX = OUT_DIR / "credit-card-mvp-spec-v9-2026-06-07.docx"
DOCX_PATH = OUT_DIR / "credit-card-mvp-spec-v10-2026-06-07.docx"
TEXT_DUMP = OUT_DIR / "credit-card-mvp-spec-v10-2026-06-07-inspection.txt"
ASSET_DIR = OUT_DIR / "assets-v20260607-v10-screens"
ASSET_DIR.mkdir(parents=True, exist_ok=True)

PALETTE = {
    "ink": (33, 43, 54),
    "navy": (20, 45, 78),
    "teal": (46, 122, 142),
    "mist": (240, 245, 247),
    "line": (207, 215, 222),
    "gold": (201, 161, 77),
    "gray": (95, 104, 112),
    "white": (255, 255, 255),
    "soft": (247, 250, 252),
    "green": (50, 130, 95),
    "red": (168, 70, 70),
}


def font(size: int, bold: bool = False):
    candidates = [
        Path("C:/Windows/Fonts/msjhbd.ttc" if bold else "C:/Windows/Fonts/msjh.ttc"),
        Path("C:/Windows/Fonts/mingliub.ttc" if bold else "C:/Windows/Fonts/mingliu.ttc"),
    ]
    for path in candidates:
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


F_TITLE = font(42, True)
F_H1 = font(30, True)
F_H2 = font(23, True)
F_BODY = font(18, False)
F_SMALL = font(16, False)
F_TINY = font(14, False)


def rounded(draw, box, fill, outline=None, radius=18, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def text(draw, xy, value, fnt=F_BODY, fill=None):
    draw.text(xy, value, font=fnt, fill=fill or PALETTE["ink"])


def title_block(draw, title, subtitle):
    text(draw, (60, 42), title, F_TITLE, PALETTE["navy"])
    text(draw, (60, 102), subtitle, F_BODY, PALETTE["ink"])
    draw.line((60, 148, 1540, 148), fill=PALETTE["line"], width=3)


def chip(draw, x, y, label, w=190):
    rounded(draw, (x, y, x + w, y + 38), PALETTE["white"], radius=11)
    text(draw, (x + 16, y + 8), label, F_SMALL, PALETTE["ink"])


def browser_frame(draw, x, y, w, h):
    rounded(draw, (x, y, x + w, y + h), PALETTE["white"], PALETTE["line"], radius=22)
    rounded(draw, (x, y, x + w, y + 54), (232, 238, 243), PALETTE["line"], radius=22)
    for i, color in enumerate([(225, 96, 86), (235, 190, 73), (92, 180, 110)]):
        draw.ellipse((x + 24 + i * 28, y + 19, x + 40 + i * 28, y + 35), fill=color)


def card_art(draw, x, y, w=148, h=98):
    rounded(draw, (x, y, x + w, y + h), (229, 236, 241), radius=16)
    text(draw, (x + 42, y + 36), "卡面圖", F_H2, PALETTE["navy"])


def create_front_home(path):
    img = Image.new("RGB", (1600, 1000), PALETTE["white"])
    d = ImageDraw.Draw(img)
    title_block(d, "前台首頁畫面：搜尋、分類入口與精選優惠", "首頁需讓使用者先搜尋、再從分類入口切入，並以後台精選與最新優惠驅動內容。")
    browser_frame(d, 60, 190, 1480, 720)
    text(d, (110, 268), "信用卡優惠查詢", F_H1, PALETTE["navy"])
    text(d, (1160, 274), "現金回饋  網購  旅遊  銀行", F_SMALL, PALETTE["gray"])
    rounded(d, (110, 340, 1080, 406), PALETTE["mist"], PALETTE["line"], radius=22)
    text(d, (140, 360), "搜尋銀行、信用卡或優惠關鍵字", F_BODY, PALETTE["gray"])
    rounded(d, (1100, 340, 1370, 406), PALETTE["teal"], radius=22)
    text(d, (1178, 360), "搜尋優惠", F_BODY, PALETTE["white"])
    cats = ["現金回饋", "網購", "旅遊", "海外", "繳費", "新戶"]
    x = 110
    for c in cats:
        rounded(d, (x, 450, x + 190, 526), PALETTE["white"], PALETTE["line"], radius=18)
        text(d, (x + 48, 474), c, F_BODY, PALETTE["navy"])
        x += 210
    text(d, (110, 586), "精選優惠", F_H2, PALETTE["navy"])
    offers = [("台新太陽卡", "一般消費最高 5%", True), ("國泰 CUBE 卡", "指定通路最高 3%", False), ("富邦 J 卡", "旅遊場景最高 5%", False)]
    x = 110
    for name, reward, tag in offers:
        rounded(d, (x, 630, x + 400, 815), PALETTE["white"], PALETTE["line"], radius=20)
        card_art(d, x + 24, 662, 130, 86)
        text(d, (x + 178, 654), name, F_H2, PALETTE["ink"])
        text(d, (x + 178, 692), reward, F_BODY, PALETTE["teal"])
        text(d, (x + 178, 728), "適合日常消費，門檻低、規則清楚。", F_SMALL, PALETTE["gray"])
        if tag:
            rounded(d, (x + 250, 764, x + 370, 800), PALETTE["gold"], radius=16)
            text(d, (x + 275, 772), "本站推薦", F_TINY, PALETTE["ink"])
        x += 430
    img.save(path)


def create_front_category(path):
    img = Image.new("RGB", (1600, 980), PALETTE["white"])
    d = ImageDraw.Draw(img)
    title_block(d, "前台分類列表頁：固定高度與二行摘要", "列表頁負責快速比較，因此推薦理由只顯示二行；完整條件由詳情頁承接。")
    text(d, (60, 198), "現金回饋", F_H1, PALETTE["navy"])
    text(d, (60, 238), "卡片高度需可控，避免少數長內容把後面卡片整片往下擠。", F_BODY, PALETTE["ink"])
    rounded(d, (60, 290, 1540, 362), PALETTE["mist"], PALETTE["line"], radius=18)
    chip(d, 84, 308, "銀行：全部", 260)
    chip(d, 364, 308, "僅看進行中", 260)
    chip(d, 644, 308, "排序：本站推薦優先", 260)
    chip(d, 924, 308, "摘要：固定二行", 260)
    rows = [
        ("1", "台新太陽卡", "台新銀行", "一般消費最高 5%", "推薦理由：適合日常消費，門檻低、規則清楚。", "若還有更多限制條件，改在詳情頁完整說明。", True),
        ("2", "國泰 CUBE 卡", "國泰世華", "指定通路最高 3%", "推薦理由：適合數位消費，主題切換邏輯明確。", "列表頁只保留比較所需重點，不延伸長文。", False),
        ("3", "富邦 J 卡", "台北富邦", "旅遊場景最高 5%", "推薦理由：旅日場景鮮明，適合作為主題整理卡片。", "完整活動細節與限制條件交由詳情頁承接。", False),
    ]
    y = 398
    for rank, name, bank, reward, line1, line2, featured in rows:
        rounded(d, (60, y, 1540, y + 148), PALETTE["white"], PALETTE["line"], radius=20)
        rounded(d, (82, y + 28, 138, y + 84), PALETTE["navy"], radius=16)
        text(d, (105, y + 35), rank, F_H2, PALETTE["white"])
        card_art(d, 166, y + 20)
        text(d, (344, y + 22), name, F_H2, PALETTE["ink"])
        text(d, (344, y + 52), bank, F_SMALL, PALETTE["gray"])
        text(d, (344, y + 82), line1, F_SMALL, PALETTE["gray"])
        text(d, (344, y + 108), line2, F_SMALL, PALETTE["gray"])
        text(d, (1060, y + 28), reward, F_H1, PALETTE["teal"])
        text(d, (1060, y + 74), "優惠期間：2026/06/01 - 2026/12/31", F_SMALL, PALETTE["ink"])
        if featured:
            rounded(d, (1328, y + 28, 1498, y + 64), PALETTE["gold"], radius=16)
            text(d, (1354, y + 36), "本站推薦", F_SMALL, PALETTE["ink"])
        y += 170
    img.save(path)


def create_front_detail(path):
    img = Image.new("RGB", (1600, 960), PALETTE["white"])
    d = ImageDraw.Draw(img)
    title_block(d, "前台優惠詳情頁：完整條件與官方來源", "詳情頁承接列表頁未展開的資訊，讓使用者在決策前看到條件、限制與來源。")
    browser_frame(d, 60, 190, 1480, 700)
    rounded(d, (110, 280, 454, 780), PALETTE["white"], PALETTE["line"], radius=22)
    card_art(d, 150, 330, 264, 170)
    text(d, (150, 548), "國泰 CUBE 卡", F_H1, PALETTE["ink"])
    text(d, (150, 592), "銀行：國泰世華", F_BODY, PALETTE["gray"])
    text(d, (150, 630), "分類：現金回饋", F_BODY, PALETTE["gray"])
    text(d, (150, 668), "狀態：進行中", F_BODY, PALETTE["teal"])
    rounded(d, (150, 716, 278, 756), PALETTE["navy"], radius=16)
    text(d, (176, 724), "指定通路", F_SMALL, PALETTE["white"])
    rounded(d, (500, 280, 1490, 780), PALETTE["white"], PALETTE["line"], radius=22)
    text(d, (540, 322), "指定通路最高 3% 回饋", F_H1, PALETTE["navy"])
    sections = [
        ("優惠期間", "2026/06/01 - 2026/12/31"),
        ("適用條件", "需於 App 切換權益方案，並符合指定通路消費條件。"),
        ("回饋上限", "每月回饋上限 300 點。"),
        ("注意事項", "若活動條件或通路名單變更，以官方公告為準。"),
        ("官方來源", "source_url 與最後校對時間需由後台維護。"),
    ]
    y = 386
    for label, value in sections:
        rounded(d, (540, y, 1450, y + 62), PALETTE["mist"], radius=16)
        text(d, (566, y + 17), label, F_BODY, PALETTE["navy"])
        text(d, (770, y + 17), value, F_SMALL, PALETTE["gray"])
        y += 76
    img.save(path)


def create_admin_list(path, title, active, columns, rows):
    img = Image.new("RGB", (1600, 1040), PALETTE["white"])
    d = ImageDraw.Draw(img)
    title_block(d, title, "後台畫面需讓管理者知道欄位用途、前台影響與可執行操作；窄螢幕時列表可轉卡片式摘要。")
    rounded(d, (60, 184, 300, 980), PALETTE["navy"], radius=26)
    menu = ["儀表板", "銀行管理", "信用卡管理", "分類管理", "優惠管理", "設定"]
    y = 238
    for item in menu:
        fill = PALETTE["gold"] if item == active else PALETTE["white"]
        color = PALETTE["ink"] if item == active else PALETTE["navy"]
        rounded(d, (82, y, 278, y + 48), fill, radius=14)
        text(d, (110, y + 11), item, F_BODY, color)
        y += 74
    rounded(d, (330, 184, 1540, 980), PALETTE["white"], PALETTE["line"], radius=24)
    text(d, (364, 220), active, F_H1, PALETTE["navy"])
    rounded(d, (1210, 218, 1490, 268), PALETTE["teal"], radius=16)
    text(d, (1260, 232), "新增 / 儲存", F_BODY, PALETTE["white"])
    table_left, table_top, table_right, row_h = 364, 304, 1508, 66
    col_w = (table_right - table_left) // len(columns)
    rounded(d, (table_left, table_top, table_right, table_top + row_h), PALETTE["mist"], PALETTE["line"], radius=14)
    for idx, col in enumerate(columns):
        text(d, (table_left + idx * col_w + 14, table_top + 22), col, F_TINY, PALETTE["navy"])
    y = table_top + row_h
    for row in rows:
        d.rectangle((table_left, y, table_right, y + row_h), fill=PALETTE["white"])
        d.line((table_left, y + row_h, table_right, y + row_h), fill=PALETTE["line"], width=1)
        for idx, value in enumerate(row):
            text(d, (table_left + idx * col_w + 14, y + 20), value, F_TINY, PALETTE["ink"])
        y += row_h
    rounded(d, (364, 870, 1508, 940), PALETTE["mist"], PALETTE["line"], radius=16)
    text(d, (390, 892), "畫面說明：欄位旁需有提示文字，新增、編輯、停用、刪除都要明確標示前台影響。", F_BODY, PALETTE["gray"])
    img.save(path)


def create_admin_offer_edit(path):
    img = Image.new("RGB", (1600, 1260), PALETTE["white"])
    d = ImageDraw.Draw(img)
    title_block(d, "後台優惠編輯畫面：欄位用途、預覽與前台影響", "單筆優惠編輯需能讓管理者看懂摘要如何產生、排序如何影響前台、來源如何校對。")
    rounded(d, (60, 184, 300, 1200), PALETTE["navy"], radius=26)
    menu = ["儀表板", "銀行管理", "信用卡管理", "分類管理", "優惠管理", "設定"]
    y = 238
    for item in menu:
        fill = PALETTE["gold"] if item == "優惠管理" else PALETTE["white"]
        color = PALETTE["ink"] if item == "優惠管理" else PALETTE["navy"]
        rounded(d, (82, y, 278, y + 48), fill, radius=14)
        text(d, (110, y + 11), item, F_BODY, color)
        y += 74
    rounded(d, (330, 184, 1540, 1200), PALETTE["white"], PALETTE["line"], radius=24)
    text(d, (364, 220), "優惠編輯：CUBE 指定通路回饋", F_H1, PALETTE["navy"])
    sections = [
        ("基本資訊", [("優惠標題", "CUBE 指定通路回饋"), ("所屬分類", "現金回饋"), ("關聯信用卡", "CUBE 卡、CUBE COMBO 卡")]),
        ("列表摘要設定", [("summary_mode", "使用系統摘要"), ("target_audience", "數位消費"), ("highlight_1", "通路多"), ("highlight_2", "活動單純"), ("summary_preview", "適合數位消費，通路多、活動單純")]),
        ("排序與上架設定", [("is_featured", "開啟"), ("recommend_score", "95"), ("sort_order", "10"), ("is_published", "開啟")]),
        ("詳情與來源", [("conditions", "需切換權益方案"), ("reward_cap", "每月 300 點"), ("source_url", "官方活動頁"), ("last_verified_at", "2026/06/07")]),
    ]
    y = 292
    for sec, fields in sections:
        h = 72 + len(fields) * 54
        rounded(d, (364, y, 1508, y + h), PALETTE["mist"], PALETTE["line"], radius=18)
        text(d, (392, y + 18), sec, F_H2, PALETTE["navy"])
        fy = y + 62
        for label, value in fields:
            rounded(d, (392, fy, 720, fy + 38), PALETTE["white"], radius=10)
            rounded(d, (744, fy, 1098, fy + 38), PALETTE["white"], radius=10)
            text(d, (410, fy + 8), label, F_SMALL, PALETTE["gray"])
            text(d, (762, fy + 8), value, F_SMALL, PALETTE["ink"])
            text(d, (1130, fy + 8), "影響前台顯示 / 排序 / SEO", F_TINY, PALETTE["teal"])
            fy += 54
        y += h + 22
    img.save(path)


def create_all_images():
    images = {
        "front_home": ASSET_DIR / "前台首頁畫面-2026-06-07-v10.png",
        "front_category": ASSET_DIR / "前台分類列表頁畫面-2026-06-07-v10.png",
        "front_detail": ASSET_DIR / "前台優惠詳情頁畫面-2026-06-07-v10.png",
        "admin_dashboard": ASSET_DIR / "後台儀表板畫面-2026-06-07-v10.png",
        "admin_bank": ASSET_DIR / "後台銀行管理畫面-2026-06-07-v10.png",
        "admin_card": ASSET_DIR / "後台信用卡管理畫面-2026-06-07-v10.png",
        "admin_category": ASSET_DIR / "後台分類管理畫面-2026-06-07-v10.png",
        "admin_offer_edit": ASSET_DIR / "後台優惠編輯畫面-2026-06-07-v10.png",
        "admin_settings": ASSET_DIR / "後台設定畫面-2026-06-07-v10.png",
    }
    create_front_home(images["front_home"])
    create_front_category(images["front_category"])
    create_front_detail(images["front_detail"])
    create_admin_list(images["admin_dashboard"], "後台儀表板畫面：統計、提醒與快捷操作", "儀表板", ["項目", "數量", "狀態", "操作"], [
        ["上架優惠", "128", "正常", "查看優惠"],
        ["即將到期", "9", "需檢查", "進入編輯"],
        ["缺卡面圖", "4", "待補", "查看信用卡"],
        ["最近更新", "12", "今日", "查看紀錄"],
    ])
    create_admin_list(images["admin_bank"], "後台銀行管理畫面：銀行資料與 Logo 維護", "銀行管理", ["銀行名稱", "slug", "Logo", "官網", "狀態", "操作"], [
        ["台新銀行", "taishin", "已上傳", "taishinbank.com.tw", "啟用", "編輯 / 停用"],
        ["國泰世華", "cathay", "已上傳", "cathaybk.com.tw", "啟用", "編輯 / 停用"],
        ["台北富邦", "fubon", "待補", "fubon.com", "停用", "編輯 / 刪除"],
    ])
    create_admin_list(images["admin_card"], "後台信用卡管理畫面：卡面圖、銀行與摘要", "信用卡管理", ["卡名", "銀行", "卡面圖", "摘要", "狀態", "操作"], [
        ["太陽卡", "台新銀行", "image_url", "日常消費", "啟用", "編輯"],
        ["CUBE 卡", "國泰世華", "image_url", "數位消費", "啟用", "編輯"],
        ["J 卡", "台北富邦", "image_url", "旅遊場景", "啟用", "編輯"],
    ])
    create_admin_list(images["admin_category"], "後台分類管理畫面：分類、排序與 SEO", "分類管理", ["分類", "slug", "icon", "排序", "SEO", "操作"], [
        ["現金回饋", "cashback", "wallet", "10", "已設定", "編輯"],
        ["網購", "online", "cart", "20", "已設定", "編輯"],
        ["旅遊", "travel", "plane", "30", "待補 FAQ", "編輯"],
    ])
    create_admin_offer_edit(images["admin_offer_edit"])
    create_admin_list(images["admin_settings"], "後台設定畫面：全站預設與顯示規則", "設定", ["欄位", "目前值", "用途", "操作"], [
        ["site_name", "信用卡優惠查詢", "Header / SEO", "儲存"],
        ["首頁精選數量", "6", "首頁卡片", "儲存"],
        ["分類頁每頁數", "12", "列表分頁", "儲存"],
        ["顯示過期優惠", "否", "前台規則", "切換"],
    ])
    return images


def style_run(run, size=11, bold=False, color="212B36"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def append_to_doc(images):
    copyfile(SRC_DOCX, DOCX_PATH)
    doc = Document(DOCX_PATH)
    doc.add_page_break()
    doc.add_paragraph("14. 前後台畫面示意圖產出規格", style="Heading 1")
    p = doc.add_paragraph()
    style_run(p.add_run("本節將前述欄位表轉成可視化畫面，用於檢查：前台每個區塊是否有後台欄位支撐，後台每個欄位是否能被管理者理解、編輯、預覽與驗收。"), 11)
    doc.add_paragraph("14.1 前台畫面示意圖", style="Heading 2")
    for key, caption in [
        ("front_home", "前台首頁：搜尋、分類入口、精選優惠與 SEO/FAQ 區塊"),
        ("front_category", "前台分類列表頁：篩選列、固定高度卡片、二行摘要與排序"),
        ("front_detail", "前台優惠詳情頁：完整條件、限制、官方來源與最後校對時間"),
    ]:
        doc.add_paragraph(caption, style="Heading 3")
        doc.add_picture(str(images[key]), width=Inches(6.35))
    doc.add_page_break()
    doc.add_paragraph("14.2 後台畫面示意圖", style="Heading 2")
    for key, caption in [
        ("admin_dashboard", "後台儀表板：統計、待處理提醒與快捷操作"),
        ("admin_bank", "銀行管理：銀行資料、Logo、官網、啟用狀態"),
        ("admin_card", "信用卡管理：卡面圖、銀行、摘要、啟用狀態"),
        ("admin_category", "分類管理：分類入口、排序、icon 與 SEO"),
        ("admin_offer_edit", "優惠編輯：摘要設定、排序上架、詳情來源與預覽"),
        ("admin_settings", "設定：全站 SEO 預設值與顯示規則"),
    ]:
        doc.add_paragraph(caption, style="Heading 3")
        doc.add_picture(str(images[key]), width=Inches(6.35))
    doc.add_paragraph("14.3 畫面產出驗收", style="Heading 2")
    for item in [
        "每張前台畫面都需能回溯到後台來源欄位。",
        "每張後台畫面都需標示主要操作與前台影響。",
        "列表卡片示意圖需呈現固定高度、二行摘要、篩選與排序。",
        "後台表單示意圖需呈現欄位值、操作方式、預覽或提示文字。",
        "新版圖檔另存於 assets-v20260607-v10-screens，不覆蓋舊圖。",
    ]:
        p = doc.add_paragraph(style="List Bullet")
        style_run(p.add_run(item), 10.5)
    doc.save(DOCX_PATH)
    TEXT_DUMP.write_text("\n".join(p.text for p in doc.paragraphs if p.text.strip()), encoding="utf-8")


if __name__ == "__main__":
    imgs = create_all_images()
    append_to_doc(imgs)
    print(DOCX_PATH)
    print(TEXT_DUMP)
    for path in imgs.values():
        print(path)
