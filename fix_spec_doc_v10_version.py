# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document


DOCX_PATH = Path(r"C:\Users\user\Documents\Credit card web project\outputs\product-spec\credit-card-mvp-spec-v10-2026-06-07.docx")
FALLBACK_DOCX_PATH = Path(r"C:\Users\user\Documents\Credit card web project\outputs\product-spec\credit-card-mvp-spec-v10-2026-06-07-fixed.docx")
TEXT_DUMP = Path(r"C:\Users\user\Documents\Credit card web project\outputs\product-spec\credit-card-mvp-spec-v10-2026-06-07-inspection.txt")

REPLACEMENTS = {
    "版本：2026-06-07 / v9": "版本：2026-06-07 / v10",
    "credit-card-mvp-spec-v9-2026-06-07.docx": "credit-card-mvp-spec-v10-2026-06-07.docx",
}


def replace_in_paragraph(paragraph):
    text = paragraph.text
    updated = text
    for old, new in REPLACEMENTS.items():
        updated = updated.replace(old, new)
    if updated == text:
        return False
    for run in paragraph.runs:
        run.text = ""
    if paragraph.runs:
        paragraph.runs[0].text = updated
    else:
        paragraph.add_run(updated)
    return True


def main():
    doc = Document(DOCX_PATH)
    changed = 0
    for paragraph in doc.paragraphs:
        if replace_in_paragraph(paragraph):
            changed += 1
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    if replace_in_paragraph(paragraph):
                        changed += 1
    saved_path = DOCX_PATH
    try:
        doc.save(DOCX_PATH)
    except PermissionError:
        saved_path = FALLBACK_DOCX_PATH
        doc.save(saved_path)
    TEXT_DUMP.write_text("\n".join(p.text for p in doc.paragraphs if p.text.strip()), encoding="utf-8")
    print(saved_path)
    print(TEXT_DUMP)
    print(f"changed={changed}")


if __name__ == "__main__":
    main()
