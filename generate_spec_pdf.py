from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont, JpegImagePlugin


ROOT = Path(__file__).resolve().parent
OUT_DIR = ROOT / "outputs" / "product-spec"
ASSET_DIR = OUT_DIR / "assets"
DOCX_PATH = OUT_DIR / "credit-card-mvp-spec.docx"
PDF_PATH = OUT_DIR / "credit-card-mvp-spec.pdf"


PALETTE = {
    "ink": (33, 43, 54),
    "navy": (20, 45, 78),
    "teal": (46, 122, 142),
    "mist": (240, 245, 247),
    "sand": (250, 246, 238),
    "gold": (201, 161, 77),
    "line": (207, 215, 222),
    "green": (49, 129, 90),
    "red": (176, 63, 69),
}


def ensure_dirs() -> None:
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def pick_font(size: int, bold: bool = False):
    candidates = [
        "C:/Windows/Fonts/msjhbd.ttc" if bold else "C:/Windows/Fonts/msjh.ttc",
        "C:/Windows/Fonts/mingliub.ttc" if bold else "C:/Windows/Fonts/mingliu.ttc",
        "C:/Windows/Fonts/arialbd.ttf" if bold else "C:/Windows/Fonts/arial.ttf",
    ]
    for candidate in candidates:
        path = Path(candidate)
        if path.exists():
            return ImageFont.truetype(str(path), size=size)
    return ImageFont.load_default()


FONT_TITLE = pick_font(52, bold=True)
FONT_H1 = pick_font(34, bold=True)
FONT_H2 = pick_font(26, bold=True)
FONT_BODY = pick_font(22)
FONT_SMALL = pick_font(18)
FONT_TINY = pick_font(16)


def rounded_box(draw: ImageDraw.ImageDraw, box, fill, outline=None, radius=24, width=2):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def draw_chip(draw, x, y, text, fill, text_fill=(255, 255, 255)):
    bbox = draw.textbbox((x, y), text, font=FONT_SMALL)
    rounded_box(draw, (x, y, bbox[2] + 32, y + 36), fill=fill, radius=18)
    draw.text((x + 16, y + 6), text, font=FONT_SMALL, fill=text_fill)


def draw_title_block(draw, title, subtitle):
    draw.text((72, 54), title, font=FONT_TITLE, fill=PALETTE["navy"])
    draw.text((72, 124), subtitle, font=FONT_BODY, fill=PALETTE["ink"])
    draw.line((72, 164, 1528, 164), fill=PALETTE["line"], width=3)


def create_homepage_mockup(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), PALETTE["sand"])
    draw = ImageDraw.Draw(img)
    draw_title_block(draw, "首頁概念：分類入口 + 精選卡片", "左側快速切分類，右側像懶人包一樣整理最新與熱門卡片。")

    rounded_box(draw, (72, 210, 370, 810), fill=(255, 255, 255), outline=PALETTE["line"])
    draw.text((102, 244), "優惠分類", font=FONT_H1, fill=PALETTE["navy"])
    categories = ["現金回饋", "里程積點", "餐飲優惠", "繳稅優惠", "旅遊優惠", "機票優惠", "更多分類"]
    y = 310
    for index, cat in enumerate(categories):
        fill = PALETTE["navy"] if index == 0 else PALETTE["mist"]
        text_fill = (255, 255, 255) if index == 0 else PALETTE["ink"]
        rounded_box(draw, (96, y, 346, y + 58), fill=fill, radius=16)
        draw.text((122, y + 13), cat, font=FONT_BODY, fill=text_fill)
        y += 74

    rounded_box(draw, (410, 210, 1528, 810), fill=(255, 255, 255), outline=PALETTE["line"])
    draw.text((448, 244), "最新信用卡 / 主題精選", font=FONT_H1, fill=PALETTE["navy"])
    draw_chip(draw, 1120, 242, "進行中優惠", PALETTE["green"])
    draw_chip(draw, 1280, 242, "精選推薦", PALETTE["gold"], text_fill=PALETTE["ink"])

    search_box = (448, 308, 1492, 368)
    rounded_box(draw, search_box, fill=PALETTE["mist"], radius=18)
    draw.text((476, 324), "搜尋卡名、銀行或優惠關鍵字", font=FONT_SMALL, fill=(104, 116, 124))

    cards = [
        ("CUBE 信用卡", "國泰世華", "現金回饋 3% 起", ["指定通路回饋", "期限清楚", "延伸優惠多"]),
        ("富邦 J 卡", "台北富邦", "旅日消費最高 5%", ["旅遊場景強", "可搭配回饋活動", "卡面辨識高"]),
        ("台新太陽卡", "台新銀行", "一般消費 2% 回饋", ["日常無腦刷", "門檻簡單", "新戶活動可延伸"]),
    ]
    top = 404
    for title, bank, reward, bullets in cards:
        rounded_box(draw, (448, top, 1492, top + 118), fill=(252, 252, 252), outline=PALETTE["line"], radius=20)
        rounded_box(draw, (476, top + 18, 646, top + 100), fill=(229, 236, 241), radius=14)
        draw.text((504, top + 42), "卡面圖", font=FONT_H2, fill=PALETTE["navy"])
        draw.text((680, top + 20), title, font=FONT_H2, fill=PALETTE["ink"])
        draw.text((680, top + 54), bank, font=FONT_SMALL, fill=(90, 102, 110))
        draw.text((944, top + 20), reward, font=FONT_H2, fill=PALETTE["teal"])
        bx = 944
        by = top + 58
        for bullet in bullets:
            draw.text((bx, by), f"• {bullet}", font=FONT_SMALL, fill=PALETTE["ink"])
            by += 24
        top += 138

    img.save(path)


def create_category_mockup(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), (251, 252, 253))
    draw = ImageDraw.Draw(img)
    draw_title_block(draw, "分類頁概念：上方篩選 + 下方卡片排序", "分類頁先解決比較問題，讓使用者一眼看到主打回饋與適用條件。")
    draw.text((72, 210), "現金回饋", font=FONT_H1, fill=PALETTE["navy"])
    draw.text((72, 254), "精選高回饋信用卡，預設依本站推薦權重排序。", font=FONT_BODY, fill=PALETTE["ink"])

    rounded_box(draw, (72, 308, 1528, 394), fill=PALETTE["mist"], outline=PALETTE["line"], radius=20)
    filters = ["銀行：全部", "關鍵字搜尋", "僅看進行中", "標籤：旅遊 / 一般消費 / 指定通路"]
    x = 104
    for label in filters:
        rounded_box(draw, (x, 330, x + 230, 372), fill=(255, 255, 255), radius=12)
        draw.text((x + 16, 341), label, font=FONT_SMALL, fill=PALETTE["ink"])
        x += 250

    items = [
        ("1", "台新信用卡", "一般消費最高 5%", "推薦原因：門檻合理、通路清楚、維護成本低", True),
        ("2", "國泰 CUBE 卡", "指定通路最高 3%", "推薦原因：通路多、主題明確、適合首頁主推", False),
        ("3", "其他卡片", "可擴充更多優惠", "保留未來多分類與多銀行資料結構", False),
    ]
    top = 438
    for rank, title, reward, note, featured in items:
        rounded_box(draw, (72, top, 1528, top + 122), fill=(255, 255, 255), outline=PALETTE["line"], radius=22)
        rounded_box(draw, (94, top + 24, 152, top + 82), fill=PALETTE["navy"], radius=18)
        draw.text((117, top + 32), rank, font=FONT_H2, fill=(255, 255, 255))
        rounded_box(draw, (182, top + 22, 322, top + 98), fill=(231, 238, 243), radius=14)
        draw.text((218, top + 48), "卡圖", font=FONT_H2, fill=PALETTE["navy"])
        draw.text((354, top + 20), title, font=FONT_H2, fill=PALETTE["ink"])
        draw.text((354, top + 58), note, font=FONT_SMALL, fill=(95, 104, 112))
        draw.text((1124, top + 28), reward, font=FONT_H1, fill=PALETTE["teal"])
        if featured:
            draw_chip(draw, 1322, top + 30, "本站推薦", PALETTE["gold"], text_fill=PALETTE["ink"])
        draw.text((1124, top + 72), "優惠期間：2026/06/01 - 2026/12/31", font=FONT_SMALL, fill=PALETTE["ink"])
        top += 144

    img.save(path)


def create_detail_mockup(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), (249, 248, 244))
    draw = ImageDraw.Draw(img)
    draw_title_block(draw, "詳情頁概念：完整條件、期間與注意事項", "列表頁解決比較，詳情頁解決決策，避免只看到高回饋卻忽略限制。")

    rounded_box(draw, (72, 214, 450, 810), fill=(255, 255, 255), outline=PALETTE["line"])
    rounded_box(draw, (106, 254, 416, 458), fill=(231, 238, 243), radius=18)
    draw.text((190, 334), "卡面圖片", font=FONT_H1, fill=PALETTE["navy"])
    draw.text((106, 498), "國泰 CUBE 信用卡", font=FONT_H2, fill=PALETTE["ink"])
    draw.text((106, 536), "分類：現金回饋", font=FONT_BODY, fill=PALETTE["ink"])
    draw.text((106, 572), "銀行：國泰世華", font=FONT_BODY, fill=PALETTE["ink"])
    draw.text((106, 608), "狀態：進行中", font=FONT_BODY, fill=PALETTE["green"])
    draw_chip(draw, 106, 654, "指定通路", PALETTE["navy"])
    draw_chip(draw, 248, 654, "數位帳單", PALETTE["teal"])

    rounded_box(draw, (490, 214, 1528, 810), fill=(255, 255, 255), outline=PALETTE["line"])
    draw.text((526, 250), "主打優惠：指定消費最高 3% 回饋", font=FONT_H1, fill=PALETTE["navy"])
    sections = [
        ("優惠期間", "2026/06/01 - 2026/12/31"),
        ("適用條件", "需切換對應權益方案，可於 App 中管理。"),
        ("回饋上限", "每月回饋上限 300 點 / 需留意活動說明。"),
        ("注意事項", "需保留原始交易條件與官方頁面連結，避免摘要與原始規則脫節。"),
        ("官方連結", "提供來源按鈕，利於後續爬蟲與人工校正。"),
    ]
    y = 326
    for label, content in sections:
        rounded_box(draw, (526, y, 1490, y + 86), fill=PALETTE["mist"], radius=18)
        draw.text((552, y + 16), label, font=FONT_H2, fill=PALETTE["ink"])
        draw.text((820, y + 22), content, font=FONT_SMALL, fill=PALETTE["ink"])
        y += 102

    img.save(path)


def create_admin_mockup(path: Path) -> None:
    img = Image.new("RGB", (1600, 900), (247, 249, 252))
    draw = ImageDraw.Draw(img)
    draw_title_block(draw, "後台概念：可維護資料，也可控排序", "MVP 的關鍵不是自動爬蟲，而是先把資料治理流程與推薦邏輯落地。")

    rounded_box(draw, (72, 212, 298, 810), fill=PALETTE["navy"], radius=28)
    menu = ["儀表板", "銀行管理", "信用卡管理", "分類管理", "優惠管理", "設定"]
    y = 262
    for item in menu:
        rounded_box(draw, (96, y, 274, y + 50), fill=(255, 255, 255, 0), outline=(255, 255, 255), radius=16)
        draw.text((124, y + 12), item, font=FONT_BODY, fill=(255, 255, 255))
        y += 76

    rounded_box(draw, (334, 212, 1528, 810), fill=(255, 255, 255), outline=PALETTE["line"], radius=24)
    draw.text((372, 248), "優惠管理", font=FONT_H1, fill=PALETTE["navy"])
    draw_chip(draw, 1304, 250, "推薦權重可手動編輯", PALETTE["gold"], text_fill=PALETTE["ink"])
    rounded_box(draw, (372, 318, 1492, 390), fill=PALETTE["mist"], radius=18)
    draw.text((398, 343), "搜尋優惠標題 / 篩銀行 / 篩分類 / 篩上架狀態", font=FONT_BODY, fill=PALETTE["ink"])

    columns = ["標題", "分類", "信用卡", "期間", "推薦分數", "上架", "操作"]
    x_positions = [386, 688, 868, 1060, 1246, 1386, 1464]
    draw.line((372, 438, 1492, 438), fill=PALETTE["line"], width=2)
    for x, col in zip(x_positions, columns):
        draw.text((x, 404), col, font=FONT_SMALL, fill=PALETTE["ink"])
    rows = [
        ("CUBE 指定通路回饋", "現金回饋", "CUBE 卡", "06/01 - 12/31", "95", "上架"),
        ("J 卡旅日回饋", "旅遊優惠", "富邦 J 卡", "07/01 - 10/31", "88", "上架"),
        ("太陽卡繳稅活動", "繳稅優惠", "太陽卡", "05/01 - 06/30", "80", "草稿"),
    ]
    y = 468
    for row in rows:
        rounded_box(draw, (372, y, 1492, y + 82), fill=(252, 252, 252), outline=PALETTE["line"], radius=16)
        values = list(row) + ["編輯"]
        for value, x in zip(values, x_positions):
            color = PALETTE["green"] if value == "上架" else PALETTE["red"] if value == "草稿" else PALETTE["ink"]
            draw.text((x, y + 25), value, font=FONT_SMALL, fill=color)
        y += 98

    img.save(path)


def set_cell_fill(cell, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    tc_pr.append(shd)


def set_page_background(section, fill_hex: str) -> None:
    sect_pr = section._sectPr
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), fill_hex)
    sect_pr.append(bg)


def style_run(run, size, bold=False, color="212B36"):
    run.font.name = "Microsoft JhengHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_title(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(title)
    style_run(run, 25, bold=True, color="14334E")
    p.space_after = Pt(4)

    p2 = doc.add_paragraph()
    run2 = p2.add_run(subtitle)
    style_run(run2, 11, color="4C5965")
    p2.space_after = Pt(12)


def add_banner_table(doc: Document, items: list[tuple[str, str]]) -> None:
    table = doc.add_table(rows=1, cols=len(items))
    table.autofit = False
    widths = [Inches(2.05), Inches(2.05), Inches(2.05)]
    for idx, width in enumerate(widths[: len(items)]):
        table.columns[idx].width = width
    for idx, (label, value) in enumerate(items):
        cell = table.cell(0, idx)
        set_cell_fill(cell, "F0F5F7")
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        r1 = p.add_run(label + "\n")
        style_run(r1, 10, bold=True, color="2E7A8E")
        r2 = p.add_run(value)
        style_run(r2, 16, bold=True, color="14334E")
    doc.add_paragraph()


def add_section_heading(doc: Document, title: str, intro: str) -> None:
    p = doc.add_paragraph()
    r = p.add_run(title)
    style_run(r, 18, bold=True, color="14334E")
    p.space_after = Pt(4)
    p2 = doc.add_paragraph()
    r2 = p2.add_run(intro)
    style_run(r2, 11, color="394651")
    p2.space_after = Pt(10)


def add_bullets(doc: Document, bullets: list[str]) -> None:
    for bullet in bullets:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(bullet)
        style_run(r, 10, color="212B36")
    doc.add_paragraph()


def add_numbered(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(style="List Number")
        r = p.add_run(item)
        style_run(r, 10, color="212B36")
    doc.add_paragraph()


def add_image_page(doc: Document, heading: str, body: str, image_path: Path) -> None:
    add_section_heading(doc, heading, body)
    doc.add_picture(str(image_path), width=Inches(6.3))
    doc.add_paragraph()


@dataclass
class TableRow:
    label: str
    value: str


def add_kv_table(doc: Document, rows: list[TableRow]) -> None:
    table = doc.add_table(rows=0, cols=2)
    table.autofit = False
    table.columns[0].width = Inches(1.9)
    table.columns[1].width = Inches(4.4)
    for row in rows:
        cells = table.add_row().cells
        set_cell_fill(cells[0], "F0F5F7")
        cells[0].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = cells[0].paragraphs[0]
        r0 = p0.add_run(row.label)
        style_run(r0, 10, bold=True, color="14334E")
        p1 = cells[1].paragraphs[0]
        r1 = p1.add_run(row.value)
        style_run(r1, 10, color="212B36")
    doc.add_paragraph()


def add_footer(section) -> None:
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = p.add_run("信用卡優惠查詢網站 MVP 規格書")
    style_run(run, 9, color="74808A")


def build_doc() -> None:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)
    add_footer(section)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft JhengHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(10.5)

    add_title(
        doc,
        "信用卡優惠查詢網站 MVP 規格書",
        "以公開查詢網站 + 後台管理為核心，先用人工維護資料建立可信任的優惠整理體驗，後續再擴充成爬蟲與規則評分。",
    )
    add_banner_table(
        doc,
        [
            ("定位", "資訊型懶人包"),
            ("版本", "MVP 第一階段"),
            ("資料策略", "人工建檔，保留爬蟲擴充"),
        ],
    )
    add_section_heading(doc, "1. 產品目標", "這個網站不是冷資料庫，而是讓使用者可以快速理解「哪張卡值得看、為什麼值得看」的入口站。")
    add_bullets(
        doc,
        [
            "前台提供公開查詢，使用者可從分類、關鍵字、銀行與是否進行中快速篩選。",
            "首頁偏導覽與整理感，透過圖文卡片、分類入口與精選區塊降低閱讀負擔。",
            "後台優先解決資料治理，讓管理者能維護銀行、信用卡、優惠、分類與推薦排序。",
        ],
    )
    add_section_heading(doc, "2. MVP 範圍", "第一版聚焦在能上線、能維護、能擴充，而不是一次做到最完整。")
    add_numbered(
        doc,
        [
            "前台：首頁、分類列表頁、優惠詳情頁、搜尋與篩選。",
            "後台：單一管理員登入，銀行 / 卡片 / 分類 / 優惠資料 CRUD。",
            "排序：採手動推薦權重，後台可編輯 is_featured、recommend_score、sort_order。",
            "資料：先人工建檔，預留未來接銀行官網爬蟲與規則評分。",
        ],
    )

    add_image_page(
        doc,
        "3. 首頁概念",
        "首頁是懶人包式入口頁。左側讓使用者快速切入優惠主題，右側集中展示最新與值得先看的信用卡。",
        ASSET_DIR / "homepage.png",
    )
    add_kv_table(
        doc,
        [
            TableRow("核心任務", "讓使用者 10 秒內找到自己想看的優惠主題。"),
            TableRow("主要元件", "分類導覽、搜尋框、精選卡片、進行中與推薦標籤。"),
            TableRow("視覺方向", "資訊清楚但不生硬，用卡面圖與重點摘要建立閱讀節奏。"),
        ],
    )

    add_image_page(
        doc,
        "4. 分類頁概念",
        "分類頁承接首頁導流，重點是比較而不是陳列全部資訊。使用者可先看主打優惠，再決定是否進一步打開詳情。",
        ASSET_DIR / "category.png",
    )
    add_kv_table(
        doc,
        [
            TableRow("預設排序", "is_featured → recommend_score → sort_order → updated_at。"),
            TableRow("列表內容", "卡面圖、卡名、銀行、主打回饋、優惠期間、推薦理由。"),
            TableRow("篩選條件", "銀行、關鍵字、是否進行中、標籤；回饋區間可作預留欄位。"),
        ],
    )

    add_image_page(
        doc,
        "5. 詳情頁概念",
        "詳情頁負責補齊條件、期限與注意事項。它不是把列表資訊重貼，而是提供決策需要的完整脈絡。",
        ASSET_DIR / "detail.png",
    )
    add_bullets(
        doc,
        [
            "卡片基本資訊：卡面、銀行、分類、狀態、標籤。",
            "主優惠資訊：主打內容、期限、門檻、回饋上限、注意事項。",
            "來源資訊：官方連結與備註，方便人工校對與後續爬蟲串接。",
        ],
    )

    add_image_page(
        doc,
        "6. 後台概念",
        "MVP 的成敗在於後台能不能穩定維護資料。這一頁不是花俏，而是要讓資料管理與推薦邏輯清楚可控。",
        ASSET_DIR / "admin.png",
    )
    add_kv_table(
        doc,
        [
            TableRow("管理模組", "儀表板、銀行管理、信用卡管理、分類管理、優惠管理、設定。"),
            TableRow("排序控制", "後台可編輯本站推薦、推薦權重、微調排序欄位。"),
            TableRow("狀態管理", "支援上架 / 草稿，避免未完成資料直接進入前台。"),
        ],
    )

    doc.add_page_break()
    add_section_heading(doc, "7. 資料結構規劃", "資料模型需同時支援前台展示與後台維護，也要為未來爬蟲與多分類延伸保留空間。")
    add_kv_table(
        doc,
        [
            TableRow("banks", "銀行名稱、slug、logo、官網、啟用狀態。"),
            TableRow("cards", "卡名、所屬銀行、卡面圖、卡片摘要、啟用狀態。"),
            TableRow("categories", "分類名稱、slug、icon、排序、啟用狀態。"),
            TableRow("offers", "標題、摘要、詳情、期間、回饋類型、回饋值、上限、條件、來源、標籤、排序控制。"),
            TableRow("offer_cards", "支援一個優惠掛多張卡，也支援一張卡有多個優惠。"),
        ],
    )
    add_section_heading(doc, "8. 商業規則與排序", "第一版不做黑盒推薦，而是用人工可控的方式建立可信任排序。")
    add_bullets(
        doc,
        [
            "前台僅顯示 is_published 的資料；若超過 end_date 可標記為已過期。",
            "推薦排序以 is_featured、recommend_score、sort_order 與 updated_at 組合。",
            "第二階段才加入規則評分機制，例如依回饋率、門檻、上限與通路廣度產生基礎分數。",
        ],
    )
    add_section_heading(doc, "9. 技術建議與開發順序", "以開發速度與後續擴充性平衡，建議前後端同專案推進。")
    add_numbered(
        doc,
        [
            "技術棧：Next.js、Prisma、SQLite、Tailwind CSS；後續可平滑升級 PostgreSQL。",
            "先完成資料庫 schema 與後台登入，再做後台 CRUD，最後串接前台首頁、分類頁與詳情頁。",
            "完成測試資料後，進行 RWD 與閱讀節奏微調，確保桌機與手機都好讀。",
        ],
    )
    add_section_heading(doc, "10. 第二階段擴充", "第一版站穩後，再把效率工具與資料自動化能力加上來。")
    add_bullets(
        doc,
        [
            "銀行官網爬蟲與半自動匯入。",
            "規則評分機制 + 人工微調推薦。",
            "SEO 強化、主題懶人包頁、更多分類與季節性專題。",
        ],
    )

    doc.save(DOCX_PATH)


def build_assets() -> None:
    create_homepage_mockup(ASSET_DIR / "homepage.png")
    create_category_mockup(ASSET_DIR / "category.png")
    create_detail_mockup(ASSET_DIR / "detail.png")
    create_admin_mockup(ASSET_DIR / "admin.png")


def wrap_lines(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = list(text)
    lines: list[str] = []
    current = ""
    for word in words:
        candidate = current + word
        if draw.textlength(candidate, font=font) <= max_width or not current:
            current = candidate
        else:
            lines.append(current)
            current = word
    if current:
        lines.append(current)
    return lines


def draw_paragraph(draw, text, x, y, width, font=FONT_BODY, fill=None, line_gap=10):
    fill = fill or PALETTE["ink"]
    lines = []
    for raw_line in text.split("\n"):
        lines.extend(wrap_lines(draw, raw_line, font, width) if raw_line else [""])
    line_height = font.size + line_gap
    for line in lines:
        draw.text((x, y), line, font=font, fill=fill)
        y += line_height
    return y


def draw_bullet_block(draw, items, x, y, width):
    for item in items:
        lines = wrap_lines(draw, item, FONT_BODY, width - 36)
        draw.text((x, y), "•", font=FONT_BODY, fill=PALETTE["teal"])
        draw.text((x + 28, y), lines[0], font=FONT_BODY, fill=PALETTE["ink"])
        line_height = FONT_BODY.size + 10
        y += line_height
        for line in lines[1:]:
            draw.text((x + 28, y), line, font=FONT_BODY, fill=PALETTE["ink"])
            y += line_height
        y += 8
    return y


def make_page_base(title: str, subtitle: str):
    img = Image.new("RGB", (1700, 2200), (250, 247, 241))
    draw = ImageDraw.Draw(img)
    draw.text((110, 100), title, font=FONT_TITLE, fill=PALETTE["navy"])
    draw.text((110, 184), subtitle, font=FONT_BODY, fill=PALETTE["ink"])
    draw.line((110, 244, 1590, 244), fill=PALETTE["line"], width=4)
    return img, draw


def card_block(draw, title, body, box, tone="mist"):
    fills = {
        "mist": (240, 245, 247),
        "white": (255, 255, 255),
        "navy": (20, 45, 78),
    }
    fill = fills[tone]
    rounded_box(draw, box, fill=fill, outline=PALETTE["line"], radius=28)
    title_color = (255, 255, 255) if tone == "navy" else PALETTE["navy"]
    body_color = (240, 245, 247) if tone == "navy" else PALETTE["ink"]
    x1, y1, x2, y2 = box
    draw.text((x1 + 28, y1 + 24), title, font=FONT_H2, fill=title_color)
    draw_paragraph(draw, body, x1 + 28, y1 + 72, x2 - x1 - 56, font=FONT_SMALL, fill=body_color, line_gap=8)


def build_pdf() -> None:
    pages = []

    img1, draw1 = make_page_base(
        "信用卡優惠查詢網站 MVP 規格書",
        "把條列規格改寫成可閱讀的提案式文件，並補上頁面概念圖與系統規劃，方便團隊快速理解方向。",
    )
    card_block(draw1, "產品定位", "公開查詢型平台，風格偏資訊整理與懶人包，不做冷資料庫感；首頁要有分類入口與精選推薦。", (110, 320, 790, 560))
    card_block(draw1, "MVP 任務", "先完成前台查詢與後台資料維護，資料採人工建檔；後續再加爬蟲與規則評分。", (830, 320, 1590, 560))
    card_block(draw1, "版本共識", "1. 公開查詢網站 + 後台管理\n2. 上方篩選、下方卡片列表\n3. 無會員，先做單一管理員登入\n4. 推薦排序採人工權重", (110, 610, 1590, 910), tone="white")
    draw1.text((110, 980), "核心原則", font=FONT_H1, fill=PALETTE["navy"])
    draw_bullet_block(
        draw1,
        [
            "先讓使用者看得懂，再追求資料自動化。",
            "首頁像整理入口，分類頁像比較頁，詳情頁像決策頁。",
            "後台要能控制本站推薦、推薦權重、上架狀態與內容來源。",
        ],
        126,
        1040,
        1460,
    )
    card_block(draw1, "使用者角色", "訪客：瀏覽、搜尋、篩選、比較。\n管理員：維護銀行、信用卡、分類、優惠與排序。", (110, 1260, 710, 1510))
    card_block(draw1, "技術建議", "Next.js + Prisma + SQLite + Tailwind CSS。MVP 用同專案完成前後台，之後可平滑升級 PostgreSQL。", (760, 1260, 1590, 1510))
    card_block(draw1, "驗收標準", "前台能依分類查優惠、後台能維護資料、分類頁能依推薦權重排序、詳情頁能看完整條件與期間。", (110, 1570, 1590, 1840), tone="navy")
    pages.append(img1)

    img2, draw2 = make_page_base("前台概念頁", "首頁與分類頁負責把大量資料壓成好讀資訊，不同頁面承擔不同決策任務。")
    home = Image.open(ASSET_DIR / "homepage.png").resize((1470, 840))
    img2.paste(home, (115, 320))
    card_block(draw2, "首頁任務", "分類導覽、搜尋入口、精選卡片。本頁負責引導與建立閱讀節奏，讓使用者不會一進站就被資料量壓垮。", (110, 1210, 760, 1500))
    card_block(draw2, "首頁元件", "左側分類、右側卡片、搜尋框、進行中標籤、推薦標記、後續可加本月熱門與即將到期。", (800, 1210, 1590, 1500))
    category = Image.open(ASSET_DIR / "category.png").resize((1470, 700))
    img2.paste(category, (115, 1530))
    pages.append(img2)

    img3, draw3 = make_page_base("詳情頁與後台概念", "列表頁解決比較，詳情頁補齊條件；後台則確保資料可控、排序可維護。")
    detail = Image.open(ASSET_DIR / "detail.png").resize((720, 405))
    admin = Image.open(ASSET_DIR / "admin.png").resize((720, 405))
    img3.paste(detail, (110, 330))
    img3.paste(admin, (870, 330))
    card_block(draw3, "詳情頁重點", "卡片基本資訊、主打優惠、期限、回饋上限、適用條件、注意事項、官方來源。讓高回饋背後的限制能被看見。", (110, 790, 760, 1080))
    card_block(draw3, "後台重點", "核心模組包含銀行、信用卡、分類、優惠管理。MVP 優先提供清楚表格與搜尋，讓人工建檔不會太痛苦。", (870, 790, 1590, 1080))
    draw3.text((110, 1160), "排序規則", font=FONT_H1, fill=PALETTE["navy"])
    draw_bullet_block(
        draw3,
        [
            "第一版以 is_featured、recommend_score、sort_order、updated_at 組合排序。",
            "不直接用高回饋百分比當排序，避免使用者被高數字但高門檻的活動誤導。",
            "第二階段才引入規則評分，讓系統自動算基礎分，再由編輯微調。",
        ],
        126,
        1220,
        1460,
    )
    card_block(draw3, "管理員流程", "登入後可查看儀表板、建立銀行與卡片資料、掛上優惠、設定期間與標籤、控制是否上架與是否推薦。", (110, 1480, 1590, 1735), tone="white")
    pages.append(img3)

    img4, draw4 = make_page_base("資料結構與內容治理", "資料模型要支援展示、查詢與後續自動化，因此在第一版就要避免把欄位做死。")
    card_block(draw4, "banks", "name / slug / logo_url / website_url / is_active / timestamps", (110, 330, 760, 520))
    card_block(draw4, "cards", "bank_id / name / slug / image_url / summary / is_active / timestamps", (830, 330, 1590, 520))
    card_block(draw4, "categories", "name / slug / icon_name / sort_order / is_active / timestamps", (110, 560, 760, 750))
    card_block(draw4, "offers", "title / summary / description / start_date / end_date / reward_type / reward_value / reward_cap / min_spend / conditions / source_url / tags / is_featured / recommend_score / sort_order / is_published", (830, 560, 1590, 940))
    card_block(draw4, "offer_cards", "offer_id / card_id，用來支援一個優惠掛多張卡，也支援一張卡同時存在多筆不同優惠。", (110, 790, 760, 980))
    draw4.text((110, 1060), "內容治理原則", font=FONT_H1, fill=PALETTE["navy"])
    draw_bullet_block(
        draw4,
        [
            "前台只顯示 is_published 的資料；end_date 過後可標成已過期。",
            "每筆優惠都保留官方來源連結，方便人工對照與未來爬蟲校正。",
            "分類與標籤分層：分類用於主導覽，標籤用於細部篩選與主題頁延伸。",
        ],
        126,
        1120,
        1460,
    )
    card_block(draw4, "為什麼先人工建檔", "信用卡優惠常有複雜條件與例外規則。第一版先建立資料治理流程，才不會在爬蟲導入後把錯誤大量自動化。", (110, 1400, 1590, 1660), tone="navy")
    pages.append(img4)

    img5, draw5 = make_page_base("開發順序與第二階段規劃", "把 MVP 做穩，比一次做太多功能更重要；第二階段再把效率與規模補上來。")
    card_block(draw5, "Phase 1：基礎架構", "建立 Next.js 專案、Prisma schema、SQLite、後台登入保護與基礎 layout。", (110, 330, 740, 560))
    card_block(draw5, "Phase 2：後台 CRUD", "先完成銀行、信用卡、分類、優惠管理，並補一批假資料驗證欄位是否足夠。", (790, 330, 1590, 560))
    card_block(draw5, "Phase 3：前台頁面", "首頁、分類頁、詳情頁，加入搜尋、篩選、推薦排序與進行中標記。", (110, 610, 740, 840))
    card_block(draw5, "Phase 4：優化與驗收", "RWD、文案節奏、卡片資訊密度、桌機 / 手機實際閱讀體驗與資料流校正。", (790, 610, 1590, 840))
    draw5.text((110, 930), "第二階段方向", font=FONT_H1, fill=PALETTE["navy"])
    draw_bullet_block(
        draw5,
        [
            "銀行官網爬蟲或半自動匯入流程。",
            "規則評分機制，先算基礎分再疊人工權重。",
            "SEO 強化、熱門主題懶人包頁、季節性活動頁。",
            "若資料量變大，再考慮 PostgreSQL、批次匯入與更完整權限管理。",
        ],
        126,
        990,
        1460,
    )
    card_block(draw5, "交付物建議", "除了這份規格 PDF，下一步可以直接銜接成 sitemap、資料表 schema、頁面元件清單與實作任務拆分。", (110, 1360, 1590, 1620), tone="white")
    pages.append(img5)

    rgb_pages = [page.convert("RGB") for page in pages]
    rgb_pages[0].save(PDF_PATH, save_all=True, append_images=rgb_pages[1:], resolution=150.0)


if __name__ == "__main__":
    ensure_dirs()
    build_assets()
    build_doc()
    build_pdf()
    print(DOCX_PATH)
    print(PDF_PATH)
