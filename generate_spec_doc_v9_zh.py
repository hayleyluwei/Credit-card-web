# -*- coding: utf-8 -*-
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(r"C:\Users\user\Documents\Credit card web project")
OUT_DIR = ROOT / "outputs" / "product-spec"
DOCX_PATH = OUT_DIR / "credit-card-mvp-spec-v9-2026-06-07.docx"
TEXT_DUMP = OUT_DIR / "credit-card-mvp-spec-v9-2026-06-07-inspection.txt"
OUT_DIR.mkdir(parents=True, exist_ok=True)

PALETTE = {
    "ink": "212B36",
    "muted": "5F6870",
    "blue": "2E74B5",
    "dark_blue": "1F4D78",
    "navy": "14334E",
    "table_fill": "E8EEF5",
    "light_fill": "F4F6F9",
    "line": "C9D3DD",
}

CONTENT = {
    "title": "信用卡優惠查詢網站 MVP 產品規格書",
    "subtitle": "版本：2026-06-07 / v9\n本版承接 v8，新增 RWD layout 規格、前台真實頁面展開、後台操作畫面與欄位可用性驗收，目標是確認欄位規劃真的能支撐後台內容維護與前台顯示。",
}


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths_dxa):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblW = tblPr.first_child_found_in("w:tblW")
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    tblW.set(qn("w:w"), str(sum(widths_dxa)))
    tblW.set(qn("w:type"), "dxa")
    tblInd = tblPr.first_child_found_in("w:tblInd")
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "120")
    tblInd.set(qn("w:type"), "dxa")
    layout = tblPr.first_child_found_in("w:tblLayout")
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tblPr.append(layout)
    layout.set(qn("w:type"), "fixed")

    old_grid = tbl.tblGrid
    if old_grid is not None:
        tbl.remove(old_grid)
    grid = OxmlElement("w:tblGrid")
    for width in widths_dxa:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    tbl.insert(1, grid)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.first_child_found_in("w:tcW")
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(widths_dxa[idx]))
            tcW.set(qn("w:type"), "dxa")
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def set_fill(cell, color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), color)


def set_borders(table, color="C9D3DD", size="4"):
    tblPr = table._tbl.tblPr
    borders = tblPr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tblPr.append(borders)
    for edge in ["top", "left", "bottom", "right", "insideH", "insideV"]:
        tag = f"w:{edge}"
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def rgb(hex_color):
    return RGBColor.from_string(hex_color)


def style_run(run, size=11, bold=False, color="212B36"):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = rgb(color)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.25):
    pf = paragraph.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = line


def configure_doc(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in [
        ("Heading 1", 16, PALETTE["blue"], 18, 10),
        ("Heading 2", 13, PALETTE["blue"], 14, 7),
        ("Heading 3", 12, PALETTE["dark_blue"], 10, 5),
    ]:
        style = doc.styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = rgb(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.line_spacing = 1.25


def add_title(doc):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 3, 1.15)
    run = p.add_run(CONTENT["title"])
    style_run(run, 24, True, PALETTE["navy"])
    p2 = doc.add_paragraph()
    set_paragraph_spacing(p2, 0, 14, 1.25)
    style_run(p2.add_run(CONTENT["subtitle"]), 10.5, False, PALETTE["muted"])


def h1(doc, text):
    doc.add_paragraph(text, style="Heading 1")


def h2(doc, text):
    doc.add_paragraph(text, style="Heading 2")


def para(doc, text, color="212B36"):
    p = doc.add_paragraph()
    set_paragraph_spacing(p, 0, 6, 1.25)
    style_run(p.add_run(text), 11, False, color)


def bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        set_paragraph_spacing(p, 0, 4, 1.25)
        style_run(p.add_run(item), 10.5, False, PALETTE["ink"])


def numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        set_paragraph_spacing(p, 0, 4, 1.25)
        style_run(p.add_run(item), 10.5, False, PALETTE["ink"])


def table(doc, headers, rows, widths):
    widths_dxa = [int(w * 1440) for w in widths]
    t = doc.add_table(rows=1, cols=len(headers))
    t.autofit = False
    set_borders(t)
    set_table_geometry(t, widths_dxa)
    for idx, header in enumerate(headers):
        cell = t.cell(0, idx)
        set_fill(cell, PALETTE["table_fill"])
        p = cell.paragraphs[0]
        set_paragraph_spacing(p, 0, 0, 1.15)
        style_run(p.add_run(header), 9.2, True, PALETTE["navy"])
    for row in rows:
        cells = t.add_row().cells
        for idx, value in enumerate(row):
            cell = cells[idx]
            if idx < len(widths_dxa):
                set_cell_margins(cell)
            p = cell.paragraphs[0]
            set_paragraph_spacing(p, 0, 0, 1.15)
            style_run(p.add_run(str(value)), 8.6, False, PALETTE["ink"])
    set_table_geometry(t, widths_dxa)
    doc.add_paragraph()
    return t


def kv_table(doc, rows):
    return table(doc, ["項目", "規格"], rows, [1.45, 5.05])


def add_callout(doc, title, body):
    t = doc.add_table(rows=1, cols=1)
    set_borders(t, "D8E0E8", "4")
    set_table_geometry(t, [9360])
    cell = t.cell(0, 0)
    set_fill(cell, PALETTE["light_fill"])
    p = cell.paragraphs[0]
    set_paragraph_spacing(p, 0, 2, 1.2)
    style_run(p.add_run(title + "："), 10.5, True, PALETTE["navy"])
    style_run(p.add_run(body), 10.5, False, PALETTE["ink"])
    doc.add_paragraph()


def build_doc():
    doc = Document()
    configure_doc(doc)
    add_title(doc)

    h1(doc, "1. 產品目標")
    para(doc, "這個網站不是冷資料庫，而是讓使用者快速理解哪張信用卡值得看、為什麼值得看的整理型入口站。")
    bullets(doc, [
        "前台提供公開查詢，使用者可從分類、關鍵字、銀行與是否進行中快速篩選。",
        "首頁偏導覽與整理感，透過圖文卡片、分類入口與精選區塊降低閱讀負擔。",
        "後台優先解決資料治理，讓管理者能維護銀行、信用卡、優惠、分類、SEO 與推薦排序。",
    ])

    h1(doc, "2. MVP 範圍")
    numbered(doc, [
        "前台：首頁、分類列表頁、搜尋結果頁、銀行頁、信用卡頁、優惠詳情頁。",
        "後台：單一管理員登入，儀表板、銀行管理、信用卡管理、分類管理、優惠管理、設定。",
        "排序：採手動推薦權重，後台可編輯 is_featured、recommend_score、sort_order。",
        "資料：先人工建檔，預留未來接銀行官網爬蟲、半自動匯入與規則評分。",
        "RWD：前台與後台需定義 mobile、tablet、desktop 三種使用狀態，並納入驗收。",
    ])

    h1(doc, "3. 技術與系統架構")
    kv_table(doc, [
        ("開發期", "Next.js / Prisma / SQLite / Tailwind CSS。先在本機完成前台、後台、CRUD、排序、摘要與 SEO 欄位驗證。"),
        ("上線期", "Next.js 部署至 Vercel 或同類型 PaaS；資料庫改為託管 PostgreSQL；持續使用 Prisma。"),
        ("圖片策略", "cards.image_url 儲存最終公開 URL；開發期可用本機或專案內測試路徑，上線期改用託管物件儲存或雲端媒體服務。"),
        ("外部來源", "銀行官網與官方優惠頁、官方申辦或導流連結；未來由排程匯入任務與爬蟲模組輔助。"),
    ])

    h1(doc, "4. 資料結構規劃")
    table(doc, ["資料實體", "核心欄位", "後台對應與用途"], [
        ("banks", "name / slug / logo_url / logo_alt / website_url / description / is_active", "銀行管理建立與維護，供信用卡關聯、銀行篩選與銀行頁使用。"),
        ("cards", "bank_id / name / slug / image_url / image_alt / summary / target_audience / is_active", "信用卡管理維護卡本身，供列表、詳情、銀行頁與信用卡頁引用。"),
        ("categories", "name / slug / icon_name / sort_order / is_active / seo_title / seo_description / faq", "分類管理維護主導覽分類、分類頁 URL、SEO 與 FAQ。"),
        ("offers", "title / category_id / summary_mode / target_audience / highlight_1 / highlight_2 / manual_summary / summary_preview / description / start_date / end_date / reward_type / reward_value / reward_cap / min_spend / conditions / source_url / last_verified_at / tags / is_featured / recommend_score / sort_order / is_published", "優惠管理建立單筆優惠，決定前台卡片、詳情頁、排序、上架與 SEO。"),
        ("offer_cards", "offer_id / card_id", "優惠與信用卡多對多中間表，支援一個優惠掛多張卡，也支援一張卡關聯多個優惠。"),
        ("settings", "site_name / default_seo_title / default_seo_description / homepage_featured_count / category_page_size / show_expired_offers", "設定頁維護全站預設值與前台顯示規則。"),
    ], [1.0, 2.75, 2.75])
    add_callout(doc, "資料關聯原則", "分類歸屬以 offers 為主，不直接掛在 cards；信用卡可透過不同優惠出現在現金回饋、網購、旅遊等多個分類中。")

    h1(doc, "5. RWD Layout 規格")
    table(doc, ["裝置", "建議寬度", "前台規則", "後台規則"], [
        ("Mobile", "320-767px", "單欄排列，搜尋與篩選優先，卡片摘要固定二行。", "側欄收合，列表改卡片式摘要，主要操作固定可見。"),
        ("Tablet", "768-1023px", "兩欄卡片，篩選可置頂或收合。", "表格保留主要欄位，次要欄位進入詳情或更多選單。"),
        ("Desktop", "1024px+", "左側分類/篩選，右側列表或精選區。", "完整側欄、表格、多欄表單與預覽區。"),
    ], [0.9, 1.1, 2.2, 2.3])
    bullets(doc, [
        "前台首頁手機版：搜尋框優先，其次分類入口，再顯示精選優惠、最新優惠與 SEO/FAQ 區塊。",
        "分類頁手機版：篩選條件預設收合；已選條件以 chip 顯示；優惠卡片一欄顯示。",
        "詳情頁手機版：卡面與主優惠先出現，條件、注意事項、官方來源依序往下。",
        "後台手機或窄螢幕：左側選單改為 drawer 或頂部選單；列表轉卡片式摘要；表單單欄排列。",
        "RWD 驗收不只看能縮放，還要檢查是否仍可搜尋、篩選、比較、編輯、預覽、儲存與上下架。",
    ])

    h1(doc, "6. 前台真實畫面規格")
    h2(doc, "6.1 首頁")
    table(doc, ["區塊", "畫面內容", "後台來源欄位", "操作 / 驗收"], [
        ("Header", "Logo、主導覽、搜尋入口", "settings.site_name、categories.name", "手機版導覽需收合。"),
        ("Hero/Search", "搜尋框、熱門分類入口", "categories、offers.tags", "可輸入關鍵字並進入搜尋結果頁。"),
        ("分類入口", "現金回饋、網購、旅遊、海外、繳費、新戶", "categories.name、slug、icon_name、sort_order", "只顯示 is_active 分類。"),
        ("精選優惠", "精選卡片列表", "offers.is_featured、recommend_score、summary_preview", "依排序規則顯示。"),
        ("最新優惠", "最近更新或上架優惠", "offers.updated_at、is_published", "僅顯示已上架資料。"),
        ("SEO/FAQ", "首頁說明與常見問題", "settings.seo、FAQ 設定", "HTML 文字可被搜尋引擎讀取。"),
    ], [1.05, 1.65, 2.0, 1.8])

    h2(doc, "6.2 分類列表頁")
    table(doc, ["區塊", "畫面內容", "後台來源欄位", "操作 / 驗收"], [
        ("分類標題", "分類名稱、分類說明", "categories.name、description、seo_title", "URL 固定，例如 /categories/cashback。"),
        ("篩選區", "銀行、是否進行中、推薦、關鍵字", "banks、offers.end_date、is_featured", "手機版可收合。"),
        ("排序", "本站推薦、分數、人工排序、更新時間", "is_featured、recommend_score、sort_order、updated_at", "與商業規則一致。"),
        ("優惠卡片", "卡面、卡名、銀行、主打回饋、期間、二行推薦理由", "cards.image_url、cards.name、banks.name、offers.summary_preview", "卡片高度穩定。"),
        ("空狀態", "無符合優惠時提示", "系統產生", "提供返回分類或清除篩選。"),
    ], [1.05, 1.7, 2.1, 1.65])

    h2(doc, "6.3 搜尋結果頁")
    table(doc, ["區塊", "畫面內容", "後台來源欄位", "操作 / 驗收"], [
        ("搜尋摘要", "顯示關鍵字與結果數", "offers、cards、banks", "搜尋結果可重複篩選。"),
        ("結果列表", "與分類頁卡片一致", "同分類頁", "不產生另一套卡片規則。"),
        ("無結果", "建議熱門分類與清除搜尋", "categories.sort_order", "不讓使用者卡死。"),
    ], [1.05, 1.7, 2.1, 1.65])

    h2(doc, "6.4 優惠詳情頁")
    table(doc, ["區塊", "畫面內容", "後台來源欄位", "操作 / 驗收"], [
        ("主資訊", "優惠標題、卡面、銀行、分類、狀態", "offers.title、cards.image_url、banks.name、categories.name", "資訊需與列表頁一致。"),
        ("主打回饋", "reward_type、reward_value、reward_cap", "offers.reward_type、reward_value、reward_cap", "數字與單位需清楚。"),
        ("條件與限制", "min_spend、conditions、注意事項", "offers.min_spend、conditions、description", "長文需有清楚段落。"),
        ("官方來源", "source_url、最後校對時間", "offers.source_url、last_verified_at", "可讓管理者追查來源。"),
        ("SEO", "title、description、canonical、JSON-LD", "offers.seo_title、seo_description、slug", "頁面需有固定 URL。"),
    ], [1.05, 1.7, 2.1, 1.65])

    h2(doc, "6.5 銀行頁")
    table(doc, ["區塊", "畫面內容", "後台來源欄位", "操作 / 驗收"], [
        ("銀行資訊", "銀行名稱、Logo、官網、介紹", "banks.name、logo_url、website_url、description", "Logo 需有 alt text。"),
        ("該銀行卡片", "信用卡列表", "cards.bank_id、is_active", "點擊可進信用卡頁。"),
        ("該銀行優惠", "優惠列表", "offer_cards、offers.is_published", "僅顯示已上架優惠。"),
    ], [1.05, 1.7, 2.1, 1.65])

    h2(doc, "6.6 信用卡頁")
    table(doc, ["區塊", "畫面內容", "後台來源欄位", "操作 / 驗收"], [
        ("卡片資訊", "卡名、銀行、卡面、摘要", "cards.name、bank_id、image_url、summary", "image_url 為公開 URL。"),
        ("適合族群", "卡片介紹與適合場景", "cards.target_audience、description", "可支援 SEO 文字。"),
        ("關聯優惠", "目前進行中優惠", "offer_cards、offers.end_date", "可區分進行中與已過期。"),
    ], [1.05, 1.7, 2.1, 1.65])

    h1(doc, "7. 後台操作畫面規格")
    h2(doc, "7.1 儀表板")
    table(doc, ["畫面區塊", "欄位 / 資料", "操作方式", "前台影響"], [
        ("統計卡片", "上架優惠、草稿、過期、卡片數、銀行數", "點擊進入對應列表", "無直接影響，用於管理判斷。"),
        ("待檢查提醒", "即將到期、來源待校對、缺圖片", "點擊進入編輯頁", "提高資料準確度。"),
        ("快捷操作", "新增優惠、新增信用卡、查看前台", "按鈕", "加速維護流程。"),
    ], [1.1, 2.0, 1.55, 1.85])

    h2(doc, "7.2 銀行管理")
    table(doc, ["欄位", "輸入方式", "必填", "預設", "前台用途", "操作提示"], [
        ("name", "文字輸入", "是", "無", "銀行名稱、篩選", "名稱應與官方一致。"),
        ("slug", "文字 / 自動產生", "是", "依 name", "銀行頁 URL", "儲存後避免任意修改。"),
        ("logo_url", "圖片 URL / 上傳", "否", "空", "銀行 Logo", "正式上線用公開 URL。"),
        ("logo_alt", "文字輸入", "否", "銀行名稱 + Logo", "圖片 alt text", "SEO 與無障礙使用。"),
        ("website_url", "URL 輸入", "否", "空", "官方連結", "需檢查 URL 格式。"),
        ("is_active", "開關", "是", "true", "是否可被選用", "停用不刪除既有關聯。"),
    ], [0.95, 1.05, 0.55, 0.9, 1.35, 1.7])

    h2(doc, "7.3 信用卡管理")
    table(doc, ["欄位", "輸入方式", "必填", "預設", "前台用途", "操作提示"], [
        ("bank_id", "下拉選單", "是", "無", "銀行與篩選", "只可選啟用銀行。"),
        ("name", "文字輸入", "是", "無", "卡名", "需能搜尋。"),
        ("slug", "文字 / 自動產生", "是", "依 name", "信用卡頁 URL", "避免重複。"),
        ("image_url", "圖片 URL / 上傳", "否", "placeholder", "卡面圖", "儲存公開可讀 URL，不存圖片本體。"),
        ("image_alt", "文字輸入", "否", "卡名 + 卡面圖", "圖片 alt text", "避免重要資訊只在圖片。"),
        ("summary", "文字輸入", "否", "空", "卡片摘要", "不取代優惠摘要。"),
        ("is_active", "開關", "是", "true", "是否可關聯優惠", "停用後不推薦新關聯。"),
    ], [0.95, 1.05, 0.55, 0.9, 1.35, 1.7])

    h2(doc, "7.4 分類管理")
    table(doc, ["欄位", "輸入方式", "必填", "預設", "前台用途", "操作提示"], [
        ("name", "文字輸入", "是", "無", "分類名稱", "例如現金回饋、網購。"),
        ("slug", "文字 / 自動產生", "是", "依 name", "分類 URL", "需唯一。"),
        ("icon_name", "下拉 / 文字", "否", "無", "分類入口 icon", "使用既有 icon set。"),
        ("sort_order", "數字輸入", "是", "0", "分類排序", "數字越小越前。"),
        ("is_active", "開關", "是", "true", "是否顯示", "停用後分類入口隱藏。"),
        ("seo_title", "文字輸入", "否", "系統產生", "SEO title", "可人工覆寫。"),
        ("seo_description", "文字輸入", "否", "系統產生", "SEO description", "建議限制長度。"),
        ("faq", "重複欄位", "否", "空", "FAQPage JSON-LD", "需與可見內容一致。"),
    ], [0.95, 1.05, 0.55, 0.9, 1.35, 1.7])

    h2(doc, "7.5 優惠管理與單筆編輯")
    table(doc, ["區塊", "欄位", "操作方式", "前台影響", "畫面說明需求"], [
        ("基本資訊", "title、category_id、offer_cards", "文字、下拉、多選", "標題、分類歸屬、關聯卡片", "說明一個優惠可關聯多張卡。"),
        ("摘要設定", "summary_mode、target_audience、highlight_1、highlight_2、manual_summary、summary_preview", "單選、下拉、文字、唯讀預覽", "列表二行推薦理由", "說明 manual_summary 有值時優先顯示。"),
        ("詳情資料", "description、conditions、reward_cap、min_spend、source_url", "長文、數字、URL", "詳情頁完整內容", "說明 source_url 用於校對。"),
        ("排序上架", "is_featured、recommend_score、sort_order、is_published、updated_at", "開關、數字、系統欄位", "是否顯示、排序位置、推薦標記", "說明排序規則。"),
        ("SEO", "seo_title、seo_description、canonical、faq", "文字 / 重複欄位", "搜尋可見性", "需與可見文字一致。"),
    ], [0.9, 1.8, 1.25, 1.3, 1.25])

    h2(doc, "7.6 設定")
    table(doc, ["欄位", "輸入方式", "用途", "驗收"], [
        ("site_name", "文字輸入", "Header、SEO 預設值", "前台可看到。"),
        ("default_seo_title", "文字輸入", "未設定頁面的預設 SEO", "metadata 有 fallback。"),
        ("default_seo_description", "文字輸入", "預設 description", "不可空白。"),
        ("homepage_featured_count", "數字輸入", "首頁精選數量", "改值後首頁數量改變。"),
        ("category_page_size", "數字輸入", "分類頁每頁筆數", "分頁或載入更多一致。"),
        ("show_expired_offers", "開關", "是否顯示過期優惠", "前台清楚標示已過期。"),
    ], [1.55, 1.25, 2.2, 1.5])

    h1(doc, "8. 前台顯示欄位 vs 後台來源欄位")
    table(doc, ["前台顯示", "後台來源", "資料表", "驗收方式"], [
        ("卡名", "信用卡管理 name", "cards", "修改後列表與詳情同步。"),
        ("銀行名稱", "銀行管理 name", "banks", "篩選與卡片顯示一致。"),
        ("卡面圖", "信用卡管理 image_url", "cards", "圖片可載入且比例穩定。"),
        ("推薦理由", "優惠管理 summary_preview 或 manual_summary", "offers", "列表最多二行。"),
        ("優惠期間", "優惠管理 start_date / end_date", "offers", "過期狀態正確。"),
        ("主打回饋", "優惠管理 reward_type / reward_value", "offers", "列表與詳情一致。"),
        ("官方來源", "優惠管理 source_url", "offers", "詳情頁可點擊。"),
        ("SEO title", "各管理模組 SEO 欄位", "categories / banks / cards / offers / settings", "metadata 正確輸出。"),
    ], [1.15, 2.1, 1.45, 1.8])

    h1(doc, "9. 商業規則與排序")
    kv_table(doc, [
        ("前台顯示原則", "前台僅顯示 is_published 的資料；若超過 end_date，可標記為已過期。"),
        ("預設排序", "is_featured → recommend_score → sort_order → updated_at。"),
        ("排序目的", "先讓編輯能控制主推內容，再用分數與微調順序處理同類優惠的前後。"),
    ])
    table(doc, ["欄位", "資料型別", "後台輸入方式", "預設值", "用途說明"], [
        ("is_featured", "Boolean", "切換開關", "false", "代表本站推薦，會影響排序與推薦標記。"),
        ("recommend_score", "Integer", "數字輸入", "0", "推薦權重分數，分數越高排序越前。"),
        ("sort_order", "Integer", "數字輸入", "0", "人工微調排序用，數字越小越前。"),
        ("updated_at", "Datetime", "系統自動帶入", "系統寫入", "其他條件相同時，以最後更新時間排序。"),
    ], [1.2, 1.0, 1.45, 0.9, 1.95])

    h1(doc, "10. SEO 與 AI 搜尋可見性")
    table(doc, ["項目", "規格"], [
        ("固定 URL", "首頁、分類頁、銀行頁、信用卡頁與優惠詳情頁都需有固定可收錄 URL，不只依賴前端篩選狀態。"),
        ("SEO metadata", "分類頁、銀行頁、信用卡頁與優惠詳情頁需可設定或產生 SEO title、SEO description、slug 與 canonical URL。"),
        ("sitemap / robots", "需建立 sitemap.xml，收錄首頁、分類頁、銀行頁、信用卡頁與優惠詳情頁；需建立 robots.txt，避免誤擋重要頁面。"),
        ("詳情頁文字內容", "優惠詳情頁需以清楚段落呈現適合族群、主打回饋、優惠期間、適用條件、回饋上限、注意事項、官方來源、最後更新時間與是否人工校對。"),
        ("結構化資料", "初期規劃 JSON-LD：WebSite、Organization、BreadcrumbList、WebPage 或 Article；若頁面有常見問題區塊，加入 FAQPage。"),
        ("圖片 alt text", "卡面圖、銀行 Logo 與示意圖需具備有意義的 alt text，避免重要資訊只存在圖片中。"),
        ("過期優惠", "過期優惠可保留頁面，但需清楚標示已過期，避免使用者與搜尋系統誤解。"),
    ], [1.35, 5.15])

    h1(doc, "11. 欄位可用性驗收清單")
    numbered(doc, [
        "這個欄位由哪個角色維護？",
        "這個欄位在哪個後台畫面新增或編輯？",
        "欄位是否必填？若不是，空值時前台怎麼顯示？",
        "欄位有沒有預設值？預設值由系統給還是管理者選？",
        "欄位是否影響前台顯示、排序、篩選、SEO 或資料匯入？",
        "欄位是否需要管理者看得懂的說明文字？",
        "欄位修改後會影響哪些前台頁面？",
        "欄位是否需要保留修改時間或最後校對時間？",
        "欄位是否可能造成資料不一致？若會，後台如何提示？",
        "手機版或窄螢幕後台是否仍能完成這個欄位的編輯？",
    ])

    h1(doc, "12. 第二階段擴充")
    bullets(doc, [
        "銀行官網爬蟲與半自動匯入。",
        "規則評分機制 + 人工微調推薦。",
        "SEO 強化、主題懶人包頁、更多分類與季節性專題。",
        "後台稽核紀錄、角色權限與內容校對工作流。",
    ])

    h1(doc, "13. 本版更新說明")
    kv_table(doc, [
        ("新增", "新增 RWD layout 規格、前台真實頁面展開、後台操作畫面展開、欄位可用性驗收清單。"),
        ("補強", "將 v8 的資料模型、SEO、圖片策略與開發期 / 上線期部署路線納入更完整的產品規格結構。"),
        ("用途", "讓後續設計、開發與驗收能檢查：每個前台畫面是否有後台欄位支撐，每個後台欄位是否真的可維護。"),
        ("版本處理", "新版檔案另存為 credit-card-mvp-spec-v9-2026-06-07.docx，未覆蓋 v8。"),
    ])

    doc.save(DOCX_PATH)
    TEXT_DUMP.write_text("\n".join(p.text for p in doc.paragraphs if p.text.strip()), encoding="utf-8")
    return DOCX_PATH, TEXT_DUMP


if __name__ == "__main__":
    docx_path, text_dump = build_doc()
    print(docx_path)
    print(text_dump)
